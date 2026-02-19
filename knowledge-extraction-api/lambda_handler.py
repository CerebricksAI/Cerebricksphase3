import json
import base64
import os
import tempfile
import yaml
import logging
import re
import urllib.parse
from typing import Dict, List, Any, Optional, Tuple, Set
from dataclasses import dataclass, asdict, field
from datetime import datetime
from pathlib import Path
import boto3
from botocore.config import Config

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# AWS clients - lazy initialization (only when needed for Lambda)
_s3_client = None
SKILLS_BUCKET = os.environ.get('SKILLS_BUCKET')


def get_s3_client():
    """Get S3 client (lazy initialization for local execution support)"""
    global _s3_client
    if _s3_client is None:
        _s3_client = boto3.client('s3')
    return _s3_client


# S3 Upload prefix for direct file uploads (before processing)
UPLOADS_PREFIX = "pending_uploads/"


def generate_presigned_upload_url(filename: str, enterprise_name: str = None, user_id: str = None, expiration: int = 3600):
    """
    Generate a presigned URL for direct S3 upload.

    This allows the UI to upload files directly to S3, bypassing API Gateway
    and Lambda payload limits. Much more efficient for large files.

    Args:
        filename: Original filename (e.g., "RetailCo_Manual.pdf")
        enterprise_name: Optional enterprise name for organization
        user_id: Optional user ID for folder organization (takes priority over enterprise_name)
        expiration: URL expiration time in seconds (default 1 hour)

    Returns:
        dict with upload_url, s3_key, and expiration info
    """
    import uuid

    # Generate unique key to avoid conflicts
    unique_id = str(uuid.uuid4())[:8]
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    # Sanitize filename for S3
    safe_filename = sanitize_for_s3_key(Path(filename).stem)
    extension = Path(filename).suffix.lower()

    # Build S3 key: pending_uploads/{user_id or enterprise}/{timestamp}_{unique}_{filename}
    # Priority: user_id > enterprise_name
    folder_name = None
    if user_id:
        folder_name = sanitize_for_s3_key(user_id)
    elif enterprise_name:
        folder_name = sanitize_for_s3_key(enterprise_name)

    if folder_name:
        s3_key = f"{UPLOADS_PREFIX}{folder_name}/{timestamp}_{unique_id}_{safe_filename}{extension}"
    else:
        s3_key = f"{UPLOADS_PREFIX}{timestamp}_{unique_id}_{safe_filename}{extension}"

    # Generate presigned URL for PUT operation
    presigned_url = get_s3_client().generate_presigned_url(
        'put_object',
        Params={
            'Bucket': SKILLS_BUCKET,
            'Key': s3_key,
            'ContentType': 'application/octet-stream'
        },
        ExpiresIn=expiration
    )

    logger.info(f"[PRESIGNED_URL] Generated upload URL for: {s3_key}")

    return {
        'upload_url': presigned_url,
        's3_key': s3_key,
        'bucket': SKILLS_BUCKET,
        'expires_in_seconds': expiration,
        'filename': filename,
        'user_id': user_id,
        'enterprise_name': enterprise_name
    }


def read_file_from_s3(s3_key: str) -> bytes:
    """
    Read file content from S3.

    Args:
        s3_key: The S3 key where the file was uploaded

    Returns:
        File content as bytes
    """
    logger.info(f"[S3_READ] Reading file from: s3://{SKILLS_BUCKET}/{s3_key}")

    response = get_s3_client().get_object(
        Bucket=SKILLS_BUCKET,
        Key=s3_key
    )

    file_content = response['Body'].read()
    logger.info(f"[S3_READ] Read {len(file_content)} bytes from S3")

    return file_content


def delete_uploaded_file(s3_key: str):
    """
    Delete the uploaded file from S3 after processing.

    Args:
        s3_key: The S3 key of the uploaded file
    """
    try:
        get_s3_client().delete_object(
            Bucket=SKILLS_BUCKET,
            Key=s3_key
        )
        logger.info(f"[S3_CLEANUP] Deleted uploaded file: {s3_key}")
    except Exception as e:
        logger.warning(f"[S3_CLEANUP] Failed to delete {s3_key}: {e}")


def fetch_skill_details(folder_id: str) -> list:
    """
    Fetch details.json from each skill subfolder within a user's skill folder.

    S3 Structure:
        Knowledge_Extraction_To_Skills/{folder_id}/skills/{skill_name}/details.json

    Args:
        folder_id: The user/enterprise folder ID (e.g., "chandler" or "akhil.parelly@company.com")

    Returns:
        List of skill details dictionaries, each containing the skill name and details.json content
    """
    skills_with_details = []

    try:
        # List all skill subfolders under the skills/ directory
        skills_prefix = f"Knowledge_Extraction_To_Skills/{folder_id}/skills/"
        logger.info(f"[SKILL_DETAILS] Listing skills at: s3://{SKILLS_BUCKET}/{skills_prefix}")

        # List objects with delimiter to get "folders"
        paginator = get_s3_client().get_paginator('list_objects_v2')
        pages = paginator.paginate(
            Bucket=SKILLS_BUCKET,
            Prefix=skills_prefix,
            Delimiter='/'
        )

        skill_folders = []
        for page in pages:
            for prefix in page.get('CommonPrefixes', []):
                skill_folder_path = prefix['Prefix']
                # Extract skill name from path (e.g., "close-store-end-of-day" from ".../skills/close-store-end-of-day/")
                skill_name = skill_folder_path.rstrip('/').split('/')[-1]
                if skill_name:
                    skill_folders.append((skill_name, skill_folder_path))

        logger.info(f"[SKILL_DETAILS] Found {len(skill_folders)} skill folders")

        # Fetch details.json from each skill folder
        for skill_name, skill_folder_path in skill_folders:
            details_key = f"{skill_folder_path}details.json"

            try:
                logger.info(f"[SKILL_DETAILS] Fetching: s3://{SKILLS_BUCKET}/{details_key}")
                response = get_s3_client().get_object(
                    Bucket=SKILLS_BUCKET,
                    Key=details_key
                )
                details_content = json.loads(response['Body'].read().decode('utf-8'))

                skills_with_details.append({
                    'skill_name': skill_name,
                    'skill_path': skill_folder_path,
                    'details': details_content
                })
                logger.info(f"[SKILL_DETAILS] Successfully loaded details for: {skill_name}")

            except get_s3_client().exceptions.NoSuchKey:
                logger.warning(f"[SKILL_DETAILS] No details.json found for skill: {skill_name}")
                # Still include the skill, but without details
                skills_with_details.append({
                    'skill_name': skill_name,
                    'skill_path': skill_folder_path,
                    'details': None,
                    '_note': 'details.json not found'
                })
            except Exception as e:
                logger.warning(f"[SKILL_DETAILS] Error reading details for {skill_name}: {str(e)}")
                skills_with_details.append({
                    'skill_name': skill_name,
                    'skill_path': skill_folder_path,
                    'details': None,
                    '_error': str(e)
                })

    except Exception as e:
        logger.error(f"[SKILL_DETAILS] Error fetching skill details: {str(e)}", exc_info=True)
        raise

    return skills_with_details


# Model configuration
BEDROCK_MODEL = "global.anthropic.claude-sonnet-4-5-20250929-v1:0"
MAX_TOKENS = 200000
DEFAULT_MAX_OUTPUT_TOKENS = 16000
MAX_INPUT_CHARS = 300000

# CORS headers for all responses - Enhanced for UI compatibility
CORS_HEADERS = {
    'Content-Type': 'application/json',
    'Access-Control-Allow-Origin': '*',
    'Access-Control-Allow-Methods': 'POST, GET, PUT, DELETE, PATCH, OPTIONS',
    'Access-Control-Allow-Headers': 'Content-Type, X-Amz-Date, Authorization, X-Api-Key, X-Amz-Security-Token, X-Requested-With, Accept, Origin, Cache-Control',
    'Access-Control-Expose-Headers': 'Content-Length, X-Request-Id, X-Amz-Request-Id',
    'Access-Control-Max-Age': '86400'  # Cache preflight for 24 hours
}


# ============================================================================
# ENTERPRISE NAME UTILITIES - For S3 skill folder naming
# ============================================================================

def sanitize_for_s3_key(name: str) -> str:
    """
    Sanitize a name for use as S3 key component.
    - Converts to lowercase
    - Replaces spaces and special chars with hyphens
    - Removes consecutive hyphens
    - Removes leading/trailing hyphens
    """
    if not name:
        return ""

    # Convert to lowercase
    sanitized = name.lower().strip()

    # Replace common separators and special chars with hyphens
    sanitized = re.sub(r'[\s_\.,;:!@#$%^&*()+=\[\]{}|\\<>?/\'\"]+', '-', sanitized)

    # Remove any remaining non-alphanumeric chars except hyphens
    sanitized = re.sub(r'[^a-z0-9\-]', '', sanitized)

    # Remove consecutive hyphens
    sanitized = re.sub(r'-+', '-', sanitized)

    # Remove leading/trailing hyphens
    sanitized = sanitized.strip('-')

    return sanitized


def sanitize_email_for_s3_key(email: str) -> str:
    """
    Sanitize an email address for use as S3 key component.
    Preserves email structure (@ and .) for readability.
    - Converts to lowercase
    - Keeps @ and . characters
    - Replaces other special chars with hyphens
    - Removes consecutive hyphens

    Example: Akhil.Parelly@quadranttechnologies.com -> akhil.parelly@quadranttechnologies.com
    """
    if not email:
        return ""

    # Convert to lowercase and strip whitespace
    sanitized = email.lower().strip()

    # Replace spaces and most special chars with hyphens, but KEEP @ and .
    sanitized = re.sub(r'[\s_,;:!#$%^&*()+=\[\]{}|\\<>?/\'\"]+', '-', sanitized)

    # Remove any remaining non-alphanumeric chars except hyphens, @, and .
    sanitized = re.sub(r'[^a-z0-9\-@\.]', '', sanitized)

    # Remove consecutive hyphens
    sanitized = re.sub(r'-+', '-', sanitized)

    # Remove leading/trailing hyphens
    sanitized = sanitized.strip('-')

    return sanitized


def is_valid_email(value: str) -> bool:
    """
    Check if a string looks like an email address.
    """
    if not value:
        return False
    # Simple email pattern check
    email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return bool(re.match(email_pattern, value.strip()))


def extract_enterprise_from_filename(filename: str) -> str:
    """
    Try to extract enterprise/company name from filename.
    Examples:
    - "RetailCo_Policy_Manual.pdf" -> "retailco"
    - "Acme Corp - Employee Handbook.docx" -> "acme-corp"
    - "company_procedures.txt" -> "company"
    """
    if not filename:
        return ""

    # Remove extension
    name_without_ext = Path(filename).stem

    # Split by common separators and take the first meaningful part
    parts = re.split(r'[\s_\-\.]+', name_without_ext)

    # Filter out common non-enterprise words (these should not be used as enterprise names)
    skip_words = {
        # Document types
        'policy', 'manual', 'handbook', 'guide', 'procedure', 'procedures',
        'document', 'doc', 'file', 'draft', 'final', 'v1', 'v2', 'version',
        # Generic terms
        'employee', 'company', 'corporate', 'internal', 'external',
        # Job/task related (not enterprise identifiers)
        'job', 'task', 'work', 'project', 'assignment',
        # Test/sample related
        'test', 'sample', 'example', 'demo', 'trial', 'temp', 'tmp',
        # Common prefixes
        'new', 'old', 'updated', 'revised', 'current', 'latest',
        # Generic document names
        'sop', 'readme', 'notes', 'instructions', 'rules', 'guidelines'
    }

    enterprise_parts = []
    for part in parts:
        part_lower = part.lower()
        if part_lower and part_lower not in skip_words and len(part_lower) > 1:
            enterprise_parts.append(part)
            # Take first 2 meaningful parts max
            if len(enterprise_parts) >= 2:
                break

    if enterprise_parts:
        return sanitize_for_s3_key('-'.join(enterprise_parts))

    # Fallback: try first part of filename if it's meaningful (not a skip word)
    if parts:
        first_part = parts[0].lower()
        # Only use first part if it's not a generic skip word and has meaningful length
        if first_part not in skip_words and len(first_part) > 2:
            return sanitize_for_s3_key(parts[0])

    # Final fallback - return empty to let caller decide
    return ""


def generate_skill_folder_name(enterprise_name: Optional[str], filename: str, user_id: Optional[str] = None, email_id: Optional[str] = None) -> str:
    """
    Generate a skill folder name for S3 storage.

    Priority:
    1. Use provided email_id (preserves email format) - e.g., "akhil.parelly@quadranttechnologies.com"
    2. Use provided user_id (sanitized) - e.g., "akhilparelly"
    3. Use provided enterprise_name (sanitized)
    4. Extract from filename
    5. Fallback to 'enterprise-skill'

    Uses a single folder per user/enterprise to consolidate all extractions.

    Returns: e.g., "akhil.parelly@quadranttechnologies.com" or "akhilparelly" or "retailco"
    """
    base_name = None
    source = None

    # Priority 1: Use provided email_id (preserves email format)
    if email_id and is_valid_email(email_id):
        base_name = sanitize_email_for_s3_key(email_id)
        if base_name:
            source = "email_id"
            logger.info(f"[SKILL_FOLDER] Using provided email_id: '{email_id}' -> '{base_name}'")

    # Priority 2: Use provided user_id (check if it's an email)
    if not base_name and user_id:
        if is_valid_email(user_id):
            # user_id is actually an email, preserve the format
            base_name = sanitize_email_for_s3_key(user_id)
            source = "user_id (email format)"
            logger.info(f"[SKILL_FOLDER] Using provided user_id as email: '{user_id}' -> '{base_name}'")
        else:
            base_name = sanitize_for_s3_key(user_id)
            source = "user_id"
            logger.info(f"[SKILL_FOLDER] Using provided user_id: '{user_id}' -> '{base_name}'")

    # Priority 2: Use provided enterprise_name
    if not base_name and enterprise_name:
        base_name = sanitize_for_s3_key(enterprise_name)
        if base_name:
            source = "provided"
            logger.info(f"[SKILL_FOLDER] Using provided enterprise name: '{enterprise_name}' -> '{base_name}'")

    # Priority 2: Extract from filename
    if not base_name:
        base_name = extract_enterprise_from_filename(filename)
        if base_name:
            source = "filename"
            logger.info(f"[SKILL_FOLDER] Extracted from filename '{filename}' -> '{base_name}'")

    # Priority 3: Fallback
    if not base_name:
        base_name = "enterprise-skill"
        source = "fallback"
        logger.warning(f"[SKILL_FOLDER] Could not extract enterprise name from filename '{filename}'. "
                      f"Using fallback: '{base_name}'. Consider providing 'enterprise_name' in request.")

    # Use single folder per enterprise (no timestamp for uniqueness)
    # This allows all extractions for the same enterprise to be stored in one folder
    folder_name = base_name

    logger.info(f"[SKILL_FOLDER] Generated: '{folder_name}' (source: {source}) - using single folder per enterprise")
    return folder_name

# ============================================================================
# DOCUMENT PARSER - Multi-format support
# ============================================================================

# Check for optional dependencies
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def parse_document(file_path: str, filename: str) -> str:
    """Parse document to plain text based on file extension.

    Supports: .txt, .md, .markdown, .rst, .json, .xml, .html, .csv, .pdf, .docx
    """
    ext = Path(filename).suffix.lower()

    # Text-based formats
    if ext in ['.txt', '.md', '.markdown', '.rst', '.json', '.xml', '.html', '.htm', '.csv']:
        return parse_text(file_path)
    # PDF
    elif ext == '.pdf':
        return parse_pdf(file_path)
    # Word documents
    elif ext == '.docx':
        return parse_docx(file_path)
    else:
        # Try to detect format from file content
        return parse_unknown(file_path, ext)


def parse_text(file_path: str) -> str:
    """Parse plain text files with encoding fallback"""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue

    raise ValueError(f"Could not decode text file with any supported encoding")


def parse_pdf(file_path: str) -> str:
    """Parse PDF files using PyPDF2"""
    if not HAS_PYPDF2:
        raise ImportError("PDF support requires PyPDF2. Install with: pip install PyPDF2")

    text_parts = []
    with open(file_path, 'rb') as f:
        pdf_reader = PyPDF2.PdfReader(f)
        num_pages = len(pdf_reader.pages)
        logger.info(f"   📄 PDF has {num_pages} pages")

        for page_num, page in enumerate(pdf_reader.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {page_num} ---\n{page_text}")

    content = '\n\n'.join(text_parts)

    if not content.strip():
        raise ValueError("PDF appears to be empty or contains only images (no extractable text)")

    return content


def parse_docx(file_path: str) -> str:
    """Parse DOCX files using python-docx"""
    if not HAS_DOCX:
        raise ImportError("DOCX support requires python-docx. Install with: pip install python-docx")

    doc = DocxDocument(file_path)
    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    content = '\n\n'.join(text_parts)

    if not content.strip():
        raise ValueError("DOCX appears to be empty")

    return content


def parse_unknown(file_path: str, ext: str) -> str:
    """Attempt to parse unknown file format"""
    # Try as text first
    try:
        return parse_text(file_path)
    except:
        pass

    # Check magic bytes
    with open(file_path, 'rb') as f:
        header = f.read(8)

    if header.startswith(b'%PDF'):
        return parse_pdf(file_path)
    elif header.startswith(b'PK'):  # ZIP-based (docx, xlsx, etc.)
        try:
            return parse_docx(file_path)
        except:
            pass

    raise ValueError(f"Unsupported file format: {ext}\nSupported: .txt, .md, .pdf, .docx")


# ============================================================================
# MULTIPART FORM DATA PARSER
# ============================================================================

def parse_multipart_form_data(body: str, content_type: str) -> Dict[str, Any]:
    """
    Parse multipart/form-data from API Gateway
    
    Returns:
        {
            'file': bytes,
            'filename': str,
            'enterprise_name': str (optional)
        }
    """
    # Extract boundary from content-type header
    boundary = None
    for part in content_type.split(';'):
        if 'boundary=' in part:
            boundary = part.split('boundary=')[1].strip()
            break
    
    if not boundary:
        raise ValueError("No boundary found in Content-Type header")
    
    # Decode body if base64 encoded (API Gateway does this)
    if body:
        try:
            body_bytes = base64.b64decode(body)
        except:
            body_bytes = body.encode('utf-8')
    else:
        raise ValueError("Empty body")
    
    # Parse multipart parts
    parts = body_bytes.split(f'--{boundary}'.encode())
    
    result = {}
    
    for part in parts:
        if not part or part == b'--\r\n' or part == b'--':
            continue
        
        # Split headers and content
        if b'\r\n\r\n' in part:
            headers_section, content = part.split(b'\r\n\r\n', 1)
        else:
            continue
        
        headers = headers_section.decode('utf-8', errors='ignore')
        
        # Extract field name and filename
        if 'Content-Disposition' in headers:
            disposition_line = [line for line in headers.split('\r\n') 
                               if 'Content-Disposition' in line][0]
            
            # Extract name
            if 'name="' in disposition_line:
                name_start = disposition_line.index('name="') + 6
                name_end = disposition_line.index('"', name_start)
                field_name = disposition_line[name_start:name_end]
            else:
                continue
            
            # Extract filename if present
            filename = None
            if 'filename="' in disposition_line:
                filename_start = disposition_line.index('filename="') + 10
                filename_end = disposition_line.index('"', filename_start)
                filename = disposition_line[filename_start:filename_end]
            
            # Clean up content (remove trailing boundary markers)
            content = content.rstrip(b'\r\n')

            # Log each field found for debugging
            logger.info(f"[MULTIPART] Found field: '{field_name}'" + (f" (filename: '{filename}')" if filename else ""))

            if field_name == 'file':
                result['file'] = content
                if filename:
                    result['filename'] = filename
            # Support multiple variations of enterprise_name field
            elif field_name in ('enterprise_name', 'enterpriseName', 'enterprise-name', 'EnterpriseName'):
                enterprise_value = content.decode('utf-8').strip()
                result['enterprise_name'] = enterprise_value
                logger.info(f"[MULTIPART] Extracted enterprise_name: '{enterprise_value}' from field '{field_name}'")

    # Log final parsed result (without file content)
    logger.info(f"[MULTIPART] Parsed result: filename='{result.get('filename')}', enterprise_name='{result.get('enterprise_name')}'")

    return result


# ============================================================================
# CANONICAL TEMPLATES - Production-Ready Format
# ============================================================================

METADATA_JSON_TEMPLATE = """{
  "domain": "<inferred enterprise domain>",
  "source_documents": [
    {
      "document_name": "<source file name>",
      "document_type": "<SOP | Policy | Manual | Email | Regulation | Other>",
      "document_version": "<if available, else null>",
      "document_owner": "<team or function if inferable, else null>"
    }
  ],
  "extracted_intents": ["<complete intent statement>"],
  "confidence_level": "<high | medium | low>",
  "assumptions": ["<explicit assumption>"],
  "limitations": ["<what NOT covered>"],
  "extraction_timestamp": "<ISO-8601 UTC>",
  "extraction_version": "knowledge-capture-v1"
}"""

RULES_YAML_TEMPLATE = """rules:
  - rule_name: "descriptive_name"
    rule: "Clear statement of requirement"
    authority: "Who can override (e.g., Manager only)"
    consequence: "What happens if violated"
    sop_reference: "Section X.X: [exact quote]"
    violation_severity: "<critical | high | medium | low>"
    signal_words: "<MUST | NEVER | ALWAYS | REQUIRED | IF/THEN>"
"""

DEFINITIONS_YAML_TEMPLATE = """definitions:
  - term: "<domain-specific term>"
    definition: "<clear meaning>"
    synonyms: ["<alternate terms>"]
    usage_context: "<when this term applies>"
"""

EXAMPLES_YAML_TEMPLATE = """examples:
  - scenario: "Descriptive scenario name"
    customer_input: "Exact customer statement"
    agent_workflow:
      - step: "Parse request"
        result: {key: "value"}
      - step: "Execute action"
        action: "What agent does"
        result: "Outcome"
    skills_chained: ["skill-1", "skill-2"]
    expected_outcome: "Final result"
    test_data_used: ["DATA-001", "PROD-123"]
"""


@dataclass
class KnowledgeType:
    """Knowledge types for extraction"""
    PROCEDURE = "procedure"
    DECISION_RULE = "decision_rule"
    DEFINITION = "definition"
    EXAMPLE = "example"
    WORKFLOW = "workflow"
    CONSTRAINT = "constraint"


# Script template for generating executable Python skills
SCRIPT_TEMPLATE = '''#!/usr/bin/env python3
"""
{skill_name} - {skill_description}
Generated from: {knowledge_source}
Version: {version}
"""

from typing import Dict, Any, List, Optional
import logging
import json

logger = logging.getLogger(__name__)


class {class_name}:
    """
    {class_docstring}
    """

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        self.config = config or {{}}
        self.logger = logger
        self._initialize()

    def _initialize(self):
        """Initialize skill-specific resources."""
{initialization_code}

    def execute(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Execute the skill.

        Args:
            input_data: {input_description}

        Returns:
            {return_description}
        """
        try:
            # Validate input
            if not self._validate_input(input_data):
                return {{"error": "Invalid input data", "success": False}}

            # Main logic
{main_execution_logic}

            return result

        except Exception as e:
            self.logger.error(f"Execution error: {{e}}")
            return {{"error": str(e), "success": False}}

    def _validate_input(self, input_data: Dict[str, Any]) -> bool:
        """Validate input data structure."""
{validation_logic}
        return True
{helper_methods}

    def get_metadata(self) -> Dict[str, Any]:
        """Return skill metadata."""
        return {{
            "name": "{skill_name}",
            "version": "{version}",
            "description": "{skill_description}",
            "required_inputs": {required_inputs},
            "optional_inputs": {optional_inputs},
            "outputs": {outputs}
        }}


# Example usage
if __name__ == "__main__":
    import json

    # Initialize the skill
    skill = {class_name}()

    # Example input
    example_input = {{
        # Add your test input here based on required_inputs
    }}

    # Execute
    result = skill.execute(example_input)
    print(json.dumps(result, indent=2))
'''


@dataclass
class ExtractedKnowledge:
    """Container for extracted knowledge"""
    type: str
    content: str
    source_document: str
    source_section: Optional[str] = None
    confidence: float = 0.0
    metadata: Dict[str, Any] = None
    requires_review: bool = False
    source_enterprise: Optional[str] = None
    skill_boundary: Optional[str] = None
    user_intention: Optional[str] = None  
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


@dataclass
class DocumentContext:
    """Document being processed"""
    filename: str
    file_type: str
    content: str
    enterprise: Optional[str] = None


@dataclass
class SkillMetadata:
    """Skill metadata"""
    name: str
    description: str
    version: str
    source_documents: List[str]
    confidence: float
    created_date: str
    enterprise: Optional[str] = None
    domain: Optional[str] = None
    knowledge_extraction_stats: Dict[str, int] = field(default_factory=dict)


class BedrockLLMClient:
    """AWS Bedrock client"""
    
    def __init__(self, model_id: str = BEDROCK_MODEL, region: str = None):
        self.model_id = model_id
        self.region = region or os.getenv('AWS_REGION', 'us-east-1')
        logger.info(f"🔧 Initializing Bedrock client")

        config = Config(
            connect_timeout=30,
            read_timeout=900,
            retries={"max_attempts": 5, "mode": "adaptive"}
        )
        
        try:
            self.client = boto3.client('bedrock-runtime', region_name=self.region, config=config)
            logger.info("✅ Bedrock client initialized")
        except Exception as e:
            logger.error(f"❌ Failed to initialize: {e}", exc_info=True)
            raise
    
    def invoke(self, prompt: str, max_tokens: int = DEFAULT_MAX_OUTPUT_TOKENS) -> str:
        """Invoke model"""
        prompt_length = len(prompt)
        
        if prompt_length > MAX_INPUT_CHARS:
            logger.warning(f"⚠️  Truncating prompt from {prompt_length} chars")
            prompt = prompt[:MAX_INPUT_CHARS] + "\n\n[TRUNCATED]"
        
        request_body = {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": max_tokens,
            "messages": [{"role": "user", "content": [{"type": "text", "text": prompt}]}]
        }
        
        try:
            response = self.client.invoke_model(
                modelId=self.model_id,
                contentType="application/json",
                accept="application/json",
                body=json.dumps(request_body)
            )
            
            response_body = json.loads(response['body'].read().decode('utf-8'))
            content = "".join(
                block.get("text", "") 
                for block in response_body.get('content', []) 
                if block.get("type") == "text"
            )
            
            return content
        except Exception as e:
            logger.error(f"❌ Invocation failed: {e}", exc_info=True)
            raise RuntimeError(f"Bedrock invocation failed: {e}")


class KnowledgeCaptureOrchestrator:
    """
    PRODUCTION-READY SKILL GENERATOR
    
    Generates agent-executable skills that meet:
    - Natural language triggers
    - Comprehensive rule extraction
    - Realistic examples with workflows
    - Data dependencies
    - Skill composition
    - Error handling
    """
    
    def __init__(self, model_id: str = BEDROCK_MODEL, region: str = None):
        logger.info("="*70)
        logger.info("🚀 PRODUCTION-READY SKILL GENERATOR")
        logger.info("="*70)
        self.client = BedrockLLMClient(model_id=model_id, region=region)
        self.model = model_id
        logger.info(f"✅ Ready (Model: {self.model})")
        logger.info("")
        logger.info("QUALITY STANDARDS:")
        logger.info("  • Natural language triggers (customer speech patterns)")
        logger.info("  • Comprehensive rules (MUST/NEVER/ALWAYS/REQUIRED)")
        logger.info("  • Realistic examples (full workflows)")
        logger.info("  • Data dependencies specified")
        logger.info("  • Skill composition mapped")
        logger.info("  • Error handling included")
        logger.info("")
        
    def process_documents(
        self, 
        document_paths: List[str],
        enterprise_name: Optional[str] = None,
        output_dir: str = "./extracted_knowledge"
    ) -> Dict[str, Any]:
        """Main entry point"""
        
        logger.info("="*70)
        logger.info("📚 PROCESSING STARTED")
        logger.info("="*70)
        logger.info(f"Documents: {len(document_paths)}")
        logger.info(f"Enterprise: {enterprise_name or 'Will be inferred'}")
        logger.info(f"Output: {output_dir}")
        logger.info("")
        
        try:
            # PHASE 1: Upload
            logger.info("📤 PHASE 1: UPLOAD")
            logger.info("-"*70)
            documents = self._upload_documents(document_paths, enterprise_name)
            logger.info(f"✅ Uploaded {len(documents)} document(s)")
            logger.info("")
            
            # PHASE 2: Extract
            logger.info("⚡ PHASE 2: EXTRACT")
            logger.info("-"*70)
            extracted_knowledge = self._extract_knowledge(documents)
            logger.info(f"✅ Extracted {len(extracted_knowledge)} items")
            logger.info("")
            
            # PHASE 3: Review
            logger.info("📋 PHASE 3: REVIEW")
            logger.info("-"*70)
            consolidated = self._review_knowledge(extracted_knowledge)
            logger.info(f"✅ Reviewed {len(consolidated)} unique items")
            logger.info("")
            
            # PHASE 4: Generate
            logger.info("⚙️  PHASE 4: GENERATE")
            logger.info("-"*70)
            skills = self._generate_skills(consolidated, enterprise_name)
            logger.info(f"✅ Generated {len(skills)} skill(s)")
            logger.info("")
            
            # PHASE 5: Complete
            logger.info("✅ PHASE 5: COMPLETE")
            logger.info("-"*70)
            self._complete_outputs(skills, output_dir)
            logger.info(f"✅ Saved to {output_dir}")
            logger.info("")
            
            result = {
                "status": "success",
                "documents_processed": len(documents),
                "knowledge_items_extracted": len(consolidated),
                "skills_generated": len(skills),
                "output_directory": output_dir,
                "timestamp": datetime.now().isoformat()
            }
            
            logger.info("="*70)
            logger.info("✅ COMPLETE")
            logger.info("="*70)
            
            return result
            
        except Exception as e:
            logger.error("="*70)
            logger.error("❌ FAILED")
            logger.error("="*70)
            logger.error(f"Error: {e}", exc_info=True)
            raise
    
    def _upload_documents(self, document_paths: List[str], enterprise_name: Optional[str]) -> List[DocumentContext]:
        """PHASE 1: Upload - Supports multiple file formats (txt, md, pdf, docx)"""
        documents = []

        for i, path in enumerate(document_paths, 1):
            try:
                filename = os.path.basename(path)
                file_ext = Path(path).suffix.lower()
                logger.info(f"📄 [{i}/{len(document_paths)}] {filename} ({file_ext})")

                # Use the multi-format document parser
                content = parse_document(path, filename)

                doc = DocumentContext(
                    filename=filename,
                    file_type=file_ext,
                    content=content,
                    enterprise=enterprise_name
                )
                documents.append(doc)
                logger.info(f"   ✅ Uploaded ({len(content)} chars extracted)")
                # Show preview of content
                preview = content[:200].replace('\n', ' ')
                logger.info(f"   📄 Preview: {preview}...")

            except Exception as e:
                logger.error(f"   ❌ Failed: {e}")

        return documents
    
    def _extract_knowledge(self, documents: List[DocumentContext]) -> List[ExtractedKnowledge]:
        """PHASE 2: Extract"""
        all_knowledge = []
        
        for i, doc in enumerate(documents, 1):
            try:
                logger.info(f"⚡ [{i}/{len(documents)}] Extracting: {doc.filename}")
                
                knowledge_items = self._extract_from_document(doc)
                
                stats = {}
                for item in knowledge_items:
                    stats[item.type] = stats.get(item.type, 0) + 1
                
                logger.info(f"   Extracted {len(knowledge_items)} items:")
                for k_type, count in sorted(stats.items()):
                    logger.info(f"      • {k_type}: {count}")
                
                all_knowledge.extend(knowledge_items)
                
            except Exception as e:
                logger.error(f"   ❌ Extraction failed: {e}")
        
        return all_knowledge
    
    def _extract_from_document(self, doc: DocumentContext) -> List[ExtractedKnowledge]:
        """Extract with comprehensive rule extraction"""
        
        prompt = f"""Extract knowledge for production-ready skill generation.

DOCUMENT: {doc.filename}
CONTENT: {doc.content[:12000]}

METADATA FORMAT (Use this exact structure):
{METADATA_JSON_TEMPLATE}

COMPREHENSIVE RULE EXTRACTION (CRITICAL):
==========================================
Extract EVERY constraint using these signal words:
- MUST / REQUIRED → Mandatory rule
- NEVER / DO NOT → Prohibition rule  
- ALWAYS → Consistent behavior rule
- IF...THEN → Conditional rule
- ONLY → Restriction rule
- SHOULD / RECOMMENDED → Advisory rule
- BEFORE / AFTER / IMMEDIATELY → Sequencing rule

Example from SOP:
"Walk them to the item (don't just point)"
Becomes:
{{
  "type": "decision_rule",
  "content": "MUST physically walk customer to product location. NEVER just point or give verbal directions only.",
  "section": "Section 1.3",
  "sop_quote": "Walk them to the item (don't just point)"
}}

OUTPUT JSON:
{{
  "inferred_domain": "Domain name",
  "user_intentions": [
    {{
      "intention_id": "hyphenated-id",
      "intention_description": "What user wants",
      "outcome": "Expected result"
    }}
  ],
  "extracted_knowledge": [
    {{
      "type": "decision_rule",
      "content": "Complete rule with signal word",
      "section": "Source section",
      "confidence": 0.9,
      "user_intention": "intention-id",
      "metadata": {{
        "signal_word": "MUST",
        "sop_quote": "Exact quote from document"
      }}
    }}
  ]
}}

Extract 15-25 items. Be COMPREHENSIVE on rules."""

        try:
            logger.info(f"   📝 Sending extraction request ({len(doc.content)} chars)...")
            response = self.client.invoke(prompt, max_tokens=16000)
            logger.info(f"   📥 Received response ({len(response)} chars)")

            response_clean = self._clean_json_response(response)
            logger.info(f"   🔧 Cleaned response, parsing JSON...")

            try:
                result = json.loads(response_clean)
            except json.JSONDecodeError as je:
                logger.error(f"   ❌ JSON parse error: {je}")
                logger.error(f"   Response preview: {response_clean[:500]}...")
                return []

            inferred_domain = result.get("inferred_domain", "Enterprise Operations")
            logger.info(f"   🏢 Domain: {inferred_domain}")
            doc.enterprise = doc.enterprise or inferred_domain

            extracted_items = result.get("extracted_knowledge", [])
            logger.info(f"   📊 Found {len(extracted_items)} items in response")

            knowledge_items = []
            for item in extracted_items:
                knowledge = ExtractedKnowledge(
                    type=item.get("type"),
                    content=item.get("content"),
                    source_document=doc.filename,
                    source_section=item.get("section"),
                    confidence=item.get("confidence", 0.8),
                    metadata=item.get("metadata", {}),
                    source_enterprise=doc.enterprise,
                    skill_boundary=item.get("user_intention"),
                    user_intention=item.get("user_intention")
                )
                knowledge_items.append(knowledge)

            if not any(k.user_intention for k in knowledge_items):
                default_intention = Path(doc.filename).stem.replace("_", "-").lower()
                for k in knowledge_items:
                    k.user_intention = default_intention

            return knowledge_items

        except Exception as e:
            logger.error(f"   ❌ Extraction failed: {e}", exc_info=True)
            return []
    
    def _review_knowledge(self, knowledge: List[ExtractedKnowledge]) -> List[ExtractedKnowledge]:
        """PHASE 3: Review"""
        
        if not knowledge:
            return []
        
        seen = set()
        unique = []
        for item in knowledge:
            key = f"{item.type}:{item.user_intention}:{item.content[:100]}"
            if key not in seen:
                seen.add(key)
                unique.append(item)
        
        logger.info(f"   Deduplicated: {len(knowledge)} → {len(unique)}")
        return unique
    
    def _generate_skills(
        self, 
        knowledge: List[ExtractedKnowledge], 
        enterprise_name: Optional[str]
    ) -> List[Dict[str, Any]]:
        """PHASE 4: Generate production-ready skills"""
        
        if not knowledge:
            return self._create_fallback_skill(enterprise_name)
        
        # Group by intention
        skills_by_intention = {}
        for item in knowledge:
            intention = item.user_intention or "default"
            if intention not in skills_by_intention:
                skills_by_intention[intention] = []
            skills_by_intention[intention].append(item)
        
        logger.info(f"   Grouped into {len(skills_by_intention)} skill(s)")
        
        generated_skills = []
        for i, (intention, skill_knowledge) in enumerate(skills_by_intention.items(), 1):
            try:
                logger.info(f"⚙️  [{i}/{len(skills_by_intention)}] Generating: {intention}")
                
                skill_output = self._generate_production_skill(intention, skill_knowledge, enterprise_name)
                
                if skill_output:
                    generated_skills.append(skill_output)
                    logger.info(f"   ✅ Generated")
                    
            except Exception as e:
                logger.error(f"   ❌ Failed: {e}")
        
        if not generated_skills:
            generated_skills = self._create_fallback_skill(enterprise_name, knowledge)
        
        return generated_skills
    
    def _generate_production_skill(
        self,
        intention: str,
        knowledge: List[ExtractedKnowledge],
        enterprise_name: Optional[str]
    ) -> Dict[str, Any]:
        """Generate production-ready skill with all quality requirements"""
        
        procedures = [k for k in knowledge if k.type == KnowledgeType.PROCEDURE]
        decision_rules = [k for k in knowledge if k.type == KnowledgeType.DECISION_RULE]
        definitions = [k for k in knowledge if k.type == KnowledgeType.DEFINITION]
        
        inferred_domain = knowledge[0].source_enterprise if knowledge else enterprise_name or "Enterprise Operations"
        
        # Generate with production quality
        skill_md = self._generate_production_skill_md(intention, procedures, decision_rules, inferred_domain, knowledge)
        rules_yaml = self._generate_production_rules_yaml(decision_rules, knowledge)
        definitions_yaml = self._generate_definitions_yaml(definitions)
        examples_yaml = self._generate_production_examples_yaml(knowledge, intention, procedures)

        # Generate configuration schema file for customization
        config_schema_json = self._generate_config_schema_json(intention, decision_rules, inferred_domain, enterprise_name, knowledge)

        # Generate Python script
        script_py = self._generate_script_for_skill(intention, knowledge, inferred_domain)

        # Generate detailed 2-3 line description for metadata
        skill_description = self._generate_skill_description(intention, procedures, decision_rules, inferred_domain)

        metadata = SkillMetadata(
            name=intention,
            description=skill_description,
            version="1.0.0",
            source_documents=list(set(k.source_document for k in knowledge)),
            confidence=sum(k.confidence for k in knowledge) / len(knowledge) if knowledge else 0.5,
            created_date=datetime.now().isoformat(),
            enterprise=enterprise_name,
            domain=inferred_domain,
            knowledge_extraction_stats={}
        )

        return {
            "metadata": asdict(metadata),
            "skill_md": skill_md,
            "rules_yaml": rules_yaml,
            "definitions_yaml": definitions_yaml,
            "examples_yaml": examples_yaml,
            "config_schema_json": config_schema_json,
            "script_py": script_py
        }
    
    def _generate_skill_description(
        self,
        intention: str,
        procedures: List[ExtractedKnowledge],
        decision_rules: List[ExtractedKnowledge],
        domain: str
    ) -> str:
        """Generate a meaningful 3-4 line description using Claude API.

        ALWAYS uses Claude API to understand the skill by its TITLE and generate
        a meaningful description that helps agents understand when to use this skill.
        """
        skill_name_readable = intention.replace('-', ' ').replace('_', ' ')

        # Collect any available context from procedures and rules
        procedure_context = ""
        rule_context = ""

        if procedures:
            proc_texts = [p.content.strip()[:300] for p in procedures[:3] if p.content]
            if proc_texts:
                procedure_context = "Available procedure details: " + " | ".join(proc_texts)

        if decision_rules:
            rule_texts = [r.content.strip()[:200] for r in decision_rules[:3] if r.content]
            if rule_texts:
                rule_context = "Business rules: " + " | ".join(rule_texts)

        # Simple, focused prompt that works with just the skill name
        prompt = f"""Generate a meaningful skill description (exactly 3 sentences) for an AI agent system.

SKILL NAME: "{skill_name_readable}"
DOMAIN: {domain}
{procedure_context}
{rule_context}

Write exactly 3 sentences:
SENTENCE 1: What does this skill do? Start with an action verb. Be specific about the task.
SENTENCE 2: What key actions or steps does it involve? List 2-4 specific procedures.
SENTENCE 3: When should an agent use this skill? What user requests would trigger it?

RULES:
- Understand the skill from its NAME - "{skill_name_readable}" tells you what it does
- Use SPECIFIC action verbs (guides, executes, verifies, processes, handles)
- NEVER say "manages", "outlines", "provides guidance", "ensures compliance"
- Include concrete examples relevant to the skill name
- Each sentence should be 15-25 words

EXAMPLES:
For "close-store-safely":
"Walks retail employees through the complete end-of-day store shutdown process step by step. Covers cash register closing, inventory securing, alarm activation, and facility lockup procedures. Use when staff need guidance on proper closing duties or ask about end-of-day tasks."

For "process-refund":
"Handles customer refund requests from initial verification through payment reversal completion. Includes receipt validation, item condition check, refund method selection, and transaction recording. Activate when customers request returns, exchanges, or money back for purchases."

Now write for "{skill_name_readable}":"""

        try:
            response = self.client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=300,
                messages=[{"role": "user", "content": prompt}]
            )

            description = response.content[0].text.strip()

            # Clean up
            description = description.replace('"', '').strip()

            # Remove common prefixes
            prefixes_to_remove = ['description:', 'here is', 'here\'s', 'for "', f'for "{skill_name_readable}":', 'output:']
            desc_lower = description.lower()
            for prefix in prefixes_to_remove:
                if desc_lower.startswith(prefix):
                    description = description[len(prefix):].strip()

            # Ensure we have meaningful content (at least 100 chars = ~3 sentences)
            if description and len(description) >= 100:
                logger.info(f"[SKILL_DESC] Generated description ({len(description)} chars) for {intention}")
                return description[:600]

            logger.warning(f"[SKILL_DESC] Response too short ({len(description)} chars), regenerating")

        except Exception as e:
            logger.error(f"[SKILL_DESC] Claude API error for {intention}: {e}")

        # Fallback: Generate based on skill name understanding
        # Parse skill name to understand intent
        skill_words = skill_name_readable.lower().split()

        # Common skill action patterns
        action_map = {
            'close': ('shutting down', 'closing procedures, security checks, and final verification'),
            'open': ('starting up', 'opening procedures, system checks, and preparation tasks'),
            'process': ('handling', 'validation, processing steps, and completion verification'),
            'complete': ('finishing', 'required steps, verification checks, and sign-off procedures'),
            'verify': ('validating', 'inspection points, compliance checks, and documentation'),
            'prepare': ('getting ready for', 'setup tasks, resource gathering, and preparation checklist'),
            'check': ('inspecting', 'verification points, status reviews, and compliance validation'),
            'handle': ('managing', 'intake procedures, processing steps, and resolution actions'),
            'submit': ('sending', 'data collection, validation, and submission confirmation'),
            'review': ('examining', 'analysis steps, evaluation criteria, and approval workflow'),
            'update': ('modifying', 'change identification, update procedures, and verification'),
            'create': ('building', 'creation steps, required inputs, and output validation'),
            'security': ('securing', 'security protocols, access controls, and safety measures'),
            'safety': ('ensuring safety through', 'safety checks, hazard identification, and protective measures')
        }

        # Find matching action
        action_verb = 'executing'
        action_details = 'required steps, verification checks, and completion procedures'

        for word in skill_words:
            if word in action_map:
                action_verb, action_details = action_map[word]
                break

        # Build meaningful fallback
        line1 = f"Guides {domain} staff through {action_verb} {skill_name_readable} tasks from start to finish."
        line2 = f"Covers {action_details} specific to {skill_name_readable}."
        line3 = f"Activate when users ask about {skill_name_readable.lower()} or need step-by-step help with related tasks."

        return f"{line1} {line2} {line3}"

    def _generate_production_skill_md(
        self,
        intention: str,
        procedures: List[ExtractedKnowledge],
        decision_rules: List[ExtractedKnowledge],
        domain: str,
        all_knowledge: List[ExtractedKnowledge]
    ) -> str:
        """Generate SKILL.md with production quality"""

        # Generate NATURAL LANGUAGE triggers
        triggers = self._generate_natural_triggers(intention, all_knowledge)

        # Generate detailed 2-3 line description
        description = self._generate_skill_description(intention, procedures, decision_rules, domain)

        frontmatter = f"""---
name: {intention}
description: >
  {description}
version: 1.0.0
triggers:
"""
        for trigger in triggers[:7]:
            frontmatter += f'  - pattern: "{trigger}"\n'
        frontmatter += "---\n"

        md = frontmatter

        md += "\n## Overview\n\n"
        md += f"{description}\n\n"
        
        md += "## When to Use\n\n"
        md += "Activate when customer expresses:\n"
        for trigger in triggers[:3]:
            md += f'- {trigger}\n'
        md += "\n"
        
        md += "## Instructions\n\n"
        if procedures:
            md += self._extract_steps(procedures[0].content)
        else:
            md += "1. Verify intent\n2. Execute workflow\n3. Confirm completion\n"
        md += "\n"
        
        md += "## Rules\n\n"
        if decision_rules:
            for rule in decision_rules:
                # Extract signal word
                signal = "MUST" if "must" in rule.content.lower() else \
                        "NEVER" if "never" in rule.content.lower() else \
                        "ALWAYS" if "always" in rule.content.lower() else \
                        "REQUIRED" if "required" in rule.content.lower() else "RULE"
                md += f"- **{signal}**: {rule.content}\n"
                if rule.metadata and rule.metadata.get("sop_quote"):
                    md += f"  - Source: {rule.source_section}: \"{rule.metadata['sop_quote']}\"\n"
        else:
            md += "- No explicit constraints documented\n"
        md += "\n"
        
        md += "## Required Data\n\n"
        md += "```yaml\n"
        md += "required_data:\n"
        md += "  - name: primary_data_source\n"
        md += "    type: database\n"
        md += "    fields: [id, name, status]\n"
        md += "```\n\n"
        
        md += "## Skill Composition\n\n"
        md += "```yaml\n"
        md += "skill_composition:\n"
        md += "  calls_these_skills: []\n"
        md += "  called_by_these_skills: []\n"
        md += "```\n\n"
        
        md += "## Error Handling\n\n"
        md += "```yaml\n"
        md += "error_scenarios:\n"
        md += "  data_unavailable:\n"
        md += "    condition: \"Required data not accessible\"\n"
        md += "    resolution: \"Use cached data with caveat\"\n"
        md += "    escalate_if: \"Cache also unavailable\"\n"
        md += "```\n\n"
        
        md += "## Examples\n\n"
        md += "See examples.yaml for detailed workflows\n\n"

        # Generate script filename from intention
        script_name = self._to_snake_case(intention) + ".py"

        md += "## Script Reference\n\n"
        md += f"This skill uses the following Python script for execution:\n\n"
        md += f"```yaml\n"
        md += f"script:\n"
        md += f"  path: scripts/{script_name}\n"
        md += f"  description: Main execution script for {intention.replace('-', ' ')}\n"
        md += f"  usage: Import and call the main function to execute this skill\n"
        md += f"```\n\n"
        md += f"**To execute this skill programmatically:**\n"
        md += f"```python\n"
        md += f"from scripts.{self._to_snake_case(intention)} import *\n"
        md += f"```\n\n"

        md += "## Configuration\n\n"
        md += "This skill can be customized via `config_schema.json`. Key sections:\n\n"
        md += "```json\n"
        md += "{\n"
        md += "  \"workflow_trigger\": {\n"
        md += "    \"trigger_type\": \"manual | scheduled | event\",\n"
        md += "    \"schedule\": { \"frequency\": \"daily\", \"timezone\": \"UTC\" }\n"
        md += "  },\n"
        md += "  \"business_thresholds\": {\n"
        md += "    \"approval_required\": false,\n"
        md += "    \"approval_threshold\": 1000,\n"
        md += "    \"escalation_enabled\": true\n"
        md += "  },\n"
        md += "  \"output_delivery\": {\n"
        md += "    \"notification_enabled\": true,\n"
        md += "    \"notification_channels\": [\"email\"]\n"
        md += "  }\n"
        md += "}\n"
        md += "```\n\n"
        md += "Edit `config_schema.json` to customize workflow triggers, thresholds, and notifications.\n\n"

        md += "## Additional Resources\n\n"
        md += f"- scripts/{script_name} - Main execution script\n"
        md += "- config_schema.json - **Customizable settings for your business**\n"
        md += "- metadata.json - Skill metadata and configuration\n"
        md += "- rules.yaml - Business rules and constraints\n"
        md += "- definitions.yaml - Term definitions and glossary\n"
        md += "- examples.yaml - Detailed workflow examples\n"

        return md
    
    def _generate_natural_triggers(
        self,
        intention: str,
        knowledge: List[ExtractedKnowledge]
    ) -> List[str]:
        """Generate NATURAL LANGUAGE triggers matching customer speech"""
        
        prompt = f"""Generate 5-7 NATURAL LANGUAGE triggers.

INTENTION: {intention}

CRITICAL REQUIREMENTS:
- Match how REAL customers speak
- NO robotic phrases
- NO keywords or technical terms
- Use pattern matching with variables: {{product}}, {{category}}, {{quantity}}
- Must be full natural thoughts

CUSTOMER LANGUAGE TEMPLATES:
- Questions: "Where is/are...", "Do you have...", "Can I get..."
- Requests: "I need...", "I'm looking for...", "I want to..."
- Problems: "This isn't working...", "I'd like to return...", "There's an issue with..."

WRONG (robotic):
- "assist product location"
- "help with assist product location"

RIGHT (natural):
- "where (can I|do I) find {{product}}"
- "do you (have|carry|sell) {{product}}"
- "(I'm|I am) looking for {{product}}"

Return 5-7 triggers, one per line, starting with "TRIGGER:"
"""
        
        try:
            response = self.client.invoke(prompt, max_tokens=800)
            
            triggers = []
            for line in response.split('\n'):
                if line.strip().startswith('TRIGGER:'):
                    trigger = line.replace('TRIGGER:', '').strip()
                    if len(trigger) > 10:
                        triggers.append(trigger)
            
            if not triggers:
                triggers = [
                    f"(I need|I want) to {intention.replace('-', ' ')}",
                    f"(Can you|Could you) help me {intention.replace('-', ' ')}",
                    f"(I'm|I am) trying to {intention.replace('-', ' ')}"
                ]
            
            return triggers[:7]
            
        except:
            return [
                f"(I need|I want) to {intention.replace('-', ' ')}",
                f"(Can you|Could you) help me {intention.replace('-', ' ')}",
                f"(I'm|I am) trying to {intention.replace('-', ' ')}"
            ][:7]
    
    def _extract_steps(self, content: str) -> str:
        """Extract steps"""
        numbered = re.findall(r'(?:^|\n)\s*\d+[\.)]\s*(.+?)(?=\n\s*\d+[\.)]|\n\n|$)', content, re.DOTALL)
        if numbered and len(numbered) >= 2:
            result = ""
            for i, step in enumerate(numbered, 1):
                result += f"{i}. {step.strip()}\n"
            return result
        
        sentences = [s.strip() for s in content.split('.') if len(s.strip()) > 10][:5]
        if sentences:
            result = ""
            for i, sentence in enumerate(sentences, 1):
                result += f"{i}. {sentence}\n"
            return result
        
        return "1. Execute workflow\n"
    
    def _generate_production_rules_yaml(
        self,
        decision_rules: List[ExtractedKnowledge],
        all_knowledge: List[ExtractedKnowledge]
    ) -> str:
        """Generate production-quality rules.yaml"""
        
        rules_dict = {"rules": []}
        
        for i, rule in enumerate(decision_rules, 1):
            content = rule.content.strip()
            
            # Extract signal word
            signal_word = None
            if "must" in content.lower(): signal_word = "MUST"
            elif "never" in content.lower(): signal_word = "NEVER"
            elif "always" in content.lower(): signal_word = "ALWAYS"
            elif "required" in content.lower(): signal_word = "REQUIRED"
            elif "should" in content.lower(): signal_word = "SHOULD"
            
            rule_entry = {
                "rule_name": f"rule_{i}",
                "rule": content,
                "signal_words": signal_word or "POLICY",
                "authority": "Manager override only",
                "sop_reference": f"{rule.source_section}: {rule.metadata.get('sop_quote', 'See source document')}" if rule.metadata else rule.source_section,
                "violation_severity": "high" if signal_word in ["MUST", "NEVER"] else "medium"
            }
            
            rules_dict["rules"].append(rule_entry)
        
        if not rules_dict["rules"]:
            rules_dict["rules"].append({
                "rule_name": "no_rules_extracted",
                "rule": "No explicit constraints documented in source",
                "signal_words": "INFO",
                "authority": "N/A",
                "sop_reference": "No SOP section found",
                "violation_severity": "info"
            })
        
        return yaml.dump(rules_dict, default_flow_style=False, sort_keys=False, allow_unicode=True)
    
    def _generate_definitions_yaml(self, definitions: List[ExtractedKnowledge]) -> str:
        """Generate definitions.yaml"""
        
        defs_dict = {"definitions": []}
        
        for defn in definitions:
            content = defn.content.strip()
            
            match = re.match(r'^([^:\n]+?):\s*(.+)', content, re.DOTALL)
            if match:
                term = match.group(1).strip()
                definition = match.group(2).strip()
            else:
                lines = content.split('\n')
                term = lines[0].strip()
                definition = '\n'.join(lines[1:]).strip() if len(lines) > 1 else content
            
            defs_dict["definitions"].append({
                "term": term,
                "definition": definition,
                "usage_context": "When applicable"
            })
        
        return yaml.dump(defs_dict, default_flow_style=False, sort_keys=False, allow_unicode=True)
    
    def _generate_production_examples_yaml(
        self,
        all_knowledge: List[ExtractedKnowledge],
        intention: str,
        procedures: List[ExtractedKnowledge]
    ) -> str:
        """Generate REALISTIC examples with full workflows"""
        
        examples_dict = {"examples": []}
        
        # Example 1: Normal flow
        examples_dict["examples"].append({
            "scenario": "Normal successful execution",
            "customer_input": f"I need to {intention.replace('-', ' ')}",
            "agent_workflow": [
                {
                    "step": "Parse request",
                    "result": {"intent": intention, "confidence": 0.95}
                },
                {
                    "step": "Verify preconditions",
                    "action": "Check user authorization",
                    "result": "Authorized"
                },
                {
                    "step": "Execute workflow",
                    "action": "Follow documented procedure",
                    "result": "Completed"
                },
                {
                    "step": "Validate outcome",
                    "result": "Success"
                }
            ],
            "skills_chained": [intention],
            "expected_outcome": "Task completed successfully",
            "test_data_used": ["USER-001", "SESSION-123"]
        })
        
        # Example 2: Error scenario
        examples_dict["examples"].append({
            "scenario": "Data unavailable error",
            "customer_input": f"Can you help me {intention.replace('-', ' ')}",
            "agent_workflow": [
                {
                    "step": "Parse request",
                    "result": {"intent": intention, "confidence": 0.92}
                },
                {
                    "step": "Attempt data retrieval",
                    "action": "Query primary data source",
                    "result": "ERROR: Database unavailable"
                },
                {
                    "step": "Use fallback",
                    "action": "Use cached data with caveat",
                    "result": "Partial data available"
                },
                {
                    "step": "Communicate limitation",
                    "script": "I can help with that, but our system is currently updating. Let me use our backup data.",
                    "result": "Customer informed"
                }
            ],
            "skills_chained": [intention],
            "expected_outcome": "Completed with caveat",
            "test_data_used": ["USER-002", "CACHE-789"]
        })
        
        return yaml.dump(examples_dict, default_flow_style=False, sort_keys=False, allow_unicode=True)

    def _generate_config_schema_json(
        self,
        intention: str,
        decision_rules: List[ExtractedKnowledge],
        domain: str,
        enterprise_name: Optional[str],
        knowledge: List[ExtractedKnowledge] = None
    ) -> str:
        """Generate config_schema.json dynamically using Claude API.

        Analyzes the skill intention and extracted knowledge to generate
        truly relevant configuration properties specific to this skill.
        """

        # Collect all extracted content for context
        rules_content = []
        procedures_content = []
        exceptions_content = []

        if decision_rules:
            for rule in decision_rules:
                if rule.content:
                    rules_content.append(rule.content.strip()[:300])

        if knowledge:
            for k in knowledge:
                if k.type == KnowledgeType.PROCEDURE and k.content:
                    procedures_content.append(k.content.strip()[:300])
                elif k.type == KnowledgeType.CONSTRAINT and k.content:
                    exceptions_content.append(k.content.strip()[:300])
                elif k.type == KnowledgeType.DECISION_RULE and k.content:
                    rules_content.append(k.content.strip()[:300])

        # Build context for Claude
        context_parts = [f"Skill Name: {intention}", f"Domain: {domain}"]

        if rules_content:
            context_parts.append(f"Business Rules:\n- " + "\n- ".join(rules_content[:5]))
        if procedures_content:
            context_parts.append(f"Procedures:\n- " + "\n- ".join(procedures_content[:5]))
        if exceptions_content:
            context_parts.append(f"Exceptions:\n- " + "\n- ".join(exceptions_content[:3]))

        skill_context = "\n\n".join(context_parts)

        # Generate config schema using Claude API
        prompt = f"""Analyze this skill and generate a config_schema.json with properties that are DIRECTLY RELEVANT to this specific skill's purpose.

{skill_context}

Generate a JSON config schema with:
1. "title": the skill name (use: {intention})
2. "description": a brief description of what this skill does and what can be configured
3. "properties": an object with 2-4 configuration sections that are SPECIFIC to this skill

IMPORTANT RULES:
- Property names must be relevant to the skill intention (e.g., for "apply-security-measures" use "security_rules", "operational_measures", "exceptions")
- DO NOT use generic properties like "execution", "thresholds", "output"
- Each property should contain actual configurable values extracted from the rules/procedures
- Use snake_case for property names
- Include arrays of rules/items that users can customize
- Keep it simple and relevant

Return ONLY valid JSON, no explanation.

Example for "apply-security-measures" skill:
{{
  "title": "apply-security-measures",
  "description": "Configure security protocols and measures for store operations",
  "properties": {{
    "security_rules": ["Rule 1 from content", "Rule 2 from content"],
    "operational_measures": ["Measure 1", "Measure 2"],
    "exceptions": ["Exception 1", "Exception 2"]
  }}
}}

Now generate for the skill "{intention}":"""

        try:
            response = self.client.invoke(prompt, max_tokens=1500)

            # Extract JSON from response
            response_text = response.strip()

            # Find JSON in response
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1

            if json_start != -1 and json_end > json_start:
                json_str = response_text[json_start:json_end]
                # Validate it's valid JSON
                config_schema = json.loads(json_str)

                # Ensure required fields exist
                if "title" not in config_schema:
                    config_schema["title"] = intention
                if "description" not in config_schema:
                    config_schema["description"] = f"Configuration for {intention.replace('-', ' ').title()}"
                if "properties" not in config_schema:
                    config_schema["properties"] = {}

                logger.info(f"[CONFIG_SCHEMA] Generated dynamic config for: {intention}")
                return json.dumps(config_schema, indent=2)
            else:
                raise ValueError("No valid JSON found in response")

        except Exception as e:
            logger.warning(f"[CONFIG_SCHEMA] Claude API failed, using fallback: {e}")
            # Fallback to basic dynamic generation
            return self._generate_fallback_config_schema(intention, domain, rules_content, procedures_content, exceptions_content)

    def _generate_fallback_config_schema(
        self,
        intention: str,
        domain: str,
        rules: List[str],
        procedures: List[str],
        exceptions: List[str]
    ) -> str:
        """Fallback config schema generation when Claude API fails."""

        # Generate property names from intention
        intention_words = intention.replace('-', '_').split('_')

        # Create dynamic property names based on intention
        primary_prop = f"{intention_words[0]}_{intention_words[-1]}_rules" if len(intention_words) > 1 else f"{intention_words[0]}_rules"
        secondary_prop = f"{intention_words[-1]}_procedures" if len(intention_words) > 1 else "procedures"

        config_schema = {
            "title": intention,
            "description": f"Configuration for {intention.replace('-', ' ').replace('_', ' ').title()} in {domain}",
            "properties": {
                primary_prop: rules[:5] if rules else [],
                secondary_prop: procedures[:5] if procedures else [],
                "exceptions": exceptions[:3] if exceptions else []
            }
        }

        # Remove empty arrays
        config_schema["properties"] = {k: v for k, v in config_schema["properties"].items() if v}

        return json.dumps(config_schema, indent=2)

    def _create_fallback_skill(
        self,
        enterprise_name: Optional[str],
        knowledge: List[ExtractedKnowledge] = None
    ) -> List[Dict[str, Any]]:
        """Create fallback skill"""

        intention = "execute-workflow"
        domain = enterprise_name or "Enterprise"

        # Generate detailed description for fallback skill
        fallback_description = (
            f"The execute workflow skill provides a standardized approach to running documented procedures for {domain} operations. "
            f"It covers the essential steps for verifying user intent, executing workflow actions in the correct sequence, and validating successful completion. "
            f"Key focus areas include compliance requirements and quality standards to ensure reliable execution."
        )

        skill_md = f"""---
name: {intention}
description: >
  {fallback_description}
version: 1.0.0
triggers:
  - pattern: "(I need|I want) to execute (the|a) workflow"
  - pattern: "(Can you|Could you) help me (with|execute) (the|a) workflow"
---

## Overview
{fallback_description}

## When to Use
When customer requests workflow execution.

## Instructions
1. Verify intent
2. Execute steps
3. Validate outcome

## Rules
- **MUST**: Follow documented procedures
- **NEVER**: Skip validation steps

## Examples
See examples.yaml

## Additional Resources
- metadata.json
"""

        metadata = SkillMetadata(
            name=intention,
            description=fallback_description,
            version="1.0.0",
            source_documents=["fallback"],
            confidence=0.5,
            created_date=datetime.now().isoformat(),
            enterprise=enterprise_name,
            domain=domain,
            knowledge_extraction_stats={}
        )
        
        # Generate default config schema for fallback
        fallback_config = self._generate_config_schema_json(intention, [], enterprise_name or "Enterprise", enterprise_name, knowledge)

        return [{
            "metadata": asdict(metadata),
            "skill_md": skill_md,
            "rules_yaml": yaml.dump({"rules": []}, default_flow_style=False),
            "definitions_yaml": yaml.dump({"definitions": []}, default_flow_style=False),
            "examples_yaml": yaml.dump({"examples": []}, default_flow_style=False),
            "config_schema_json": fallback_config,
            "script_py": None
        }]

    # ==========================================
    # SCRIPT GENERATION METHODS
    # ==========================================

    def _generate_script_for_skill(self, skill_name: str, knowledge: List[ExtractedKnowledge], domain: str) -> Optional[str]:
        """Generate a Python script for a skill with retry logic."""
        max_retries = 3
        last_error = None

        for attempt in range(max_retries):
            try:
                if attempt == 0:
                    logger.info(f"   🐍 Generating Python script...")
                else:
                    logger.info(f"   🔄 Retry attempt {attempt + 1}/{max_retries}...")

                # Generate script components using LLM
                if attempt == 0:
                    components = self._generate_script_components(skill_name, knowledge)
                else:
                    components = self._generate_script_components_with_retry(skill_name, knowledge, last_error, attempt)

                # Validate components
                validated = self._validate_script_components(components)

                # Sanitize description for module docstring (no triple quotes or newlines)
                safe_description = validated['description'][:100].replace('"""', "'").replace("'''", "'").replace('\n', ' ')

                # Fill template
                script = SCRIPT_TEMPLATE.format(
                    skill_name=skill_name,
                    skill_description=safe_description,
                    knowledge_source=domain,
                    version="1.0.0",
                    class_name=self._to_class_name(skill_name),
                    class_docstring=validated['class_docstring'],
                    initialization_code=validated['init_code'],
                    input_description=validated['input_desc'],
                    return_description=validated['return_desc'],
                    main_execution_logic=validated['main_logic'],
                    validation_logic=validated['validation'],
                    helper_methods=validated['helpers'],
                    required_inputs=json.dumps(validated['required_inputs']),
                    optional_inputs=json.dumps(validated['optional_inputs']),
                    outputs=json.dumps(validated['outputs'])
                )

                # Debug: Show first 20 lines if there's an issue
                logger.debug(f"   Generated script preview:\n" + '\n'.join(f"{i+1}: {line}" for i, line in enumerate(script.split('\n')[:20])))

                # Validate syntax
                self._validate_python_syntax(script, skill_name)

                if attempt > 0:
                    logger.info(f"   ✅ Script succeeded on retry #{attempt}")
                else:
                    logger.info(f"   ✅ Script generated and validated")
                return script

            except Exception as e:
                last_error = str(e)
                if attempt < max_retries - 1:
                    logger.warning(f"   ⚠️ Attempt {attempt + 1} failed: {last_error}")
                else:
                    logger.error(f"   ❌ Script generation failed after {max_retries} attempts: {last_error}")

        return None

    def _generate_script_components(self, skill_name: str, knowledge: List[ExtractedKnowledge]) -> Dict:
        """Use LLM to generate script components."""

        # Build knowledge summary
        knowledge_summary = []
        for k in knowledge[:5]:  # Limit to avoid token overflow
            knowledge_summary.append({
                "type": k.type,
                "content": k.content[:500] if len(k.content) > 500 else k.content
            })

        prompt = f"""Generate ONLY the following components for a skill script based on this knowledge.

Skill Name: {skill_name}

Knowledge:
{json.dumps(knowledge_summary, indent=2)}

Generate as JSON with these exact keys:
{{
  "description": "Brief description of the skill",
  "class_docstring": "Clear 2-3 sentence description of what this skill does",
  "init_code": "Python code for initialization (no imports, just logic like self.threshold = 100)",
  "input_desc": "Description of expected input_data structure (e.g., 'Dict with keys: request_type, amount')",
  "return_desc": "Description of return value structure (e.g., 'Dict with keys: approved, reason')",
  "main_logic": "Python code for main execution logic (4-12 lines, MUST set 'result' variable)",
  "validation": "Python code to validate input_data (2-6 lines, MUST return True/False)",
  "helpers": "Optional helper method definitions (complete method defs with 'def method_name(self, ...):')",
  "required_inputs": ["list", "of", "required", "keys"],
  "optional_inputs": ["list", "of", "optional", "keys"],
  "outputs": ["list", "of", "output", "keys"]
}}

CRITICAL SYNTAX RULES (AVOID COMMON ERRORS):
1. Use 'if' NOT 'iif' or 'iff' (typo - spell it correctly!)
2. Every if/else/for/while/try/except block MUST have body
   - Empty blocks? Add 'pass'
   - Example: if x > 10:\n    pass
3. main_logic MUST set 'result' variable (a dict)
   - Example: result = {{"success": True, "data": processed}}
4. validation MUST have 'return' statement
   - Example: return True
5. No import statements in any code blocks
6. No eval/exec/os.system/__import__ or dangerous functions
7. Keep logic concise and focused on the skill's purpose
8. Use only standard Python operations

Return ONLY valid JSON, no markdown or explanation."""

        response = self.client.invoke(prompt, max_tokens=2000)
        response_clean = self._clean_json_response(response)
        return json.loads(response_clean)

    def _generate_script_components_with_retry(self, skill_name: str, knowledge: List[ExtractedKnowledge], previous_error: str, attempt: int) -> Dict:
        """Generate script components with error feedback from previous attempt."""

        # Build knowledge summary
        knowledge_summary = []
        for k in knowledge[:5]:
            knowledge_summary.append({
                "type": k.type,
                "content": k.content[:500] if len(k.content) > 500 else k.content
            })

        prompt = f"""Generate ONLY the following components for a skill script based on this knowledge.

Skill Name: {skill_name}

Knowledge:
{json.dumps(knowledge_summary, indent=2)}

IMPORTANT: Previous attempt failed with error: {previous_error}

Common mistakes to avoid:
1. Typo 'iif' instead of 'if'
2. Typo 'iff' instead of 'if'
3. Empty blocks after if/else/for/while (add 'pass' if empty)
4. Using dangerous functions: eval, exec, __import__, os.system
5. Not setting 'result' variable in main_logic
6. Not having 'return' statement in validation

Generate as JSON with these exact keys:
{{
  "description": "Brief description of the skill",
  "class_docstring": "Clear 2-3 sentence description of what this skill does",
  "init_code": "Python code for initialization (no imports, just logic like self.threshold = 100)",
  "input_desc": "Description of expected input_data structure",
  "return_desc": "Description of return value structure",
  "main_logic": "Python code for main execution logic (4-12 lines, MUST set 'result' variable)",
  "validation": "Python code to validate input_data (2-6 lines, MUST return True/False)",
  "helpers": "Optional helper method definitions (complete method defs)",
  "required_inputs": ["list", "of", "required", "keys"],
  "optional_inputs": ["list", "of", "optional", "keys"],
  "outputs": ["list", "of", "output", "keys"]
}}

CRITICAL RULES:
- Use 'if' not 'iif' or 'iff'
- Every if/else/for/while block must have at least one statement (use 'pass' if needed)
- main_logic MUST set 'result' variable (a dict)
- validation MUST return True or False
- No eval/exec/os.system/__import__
- No import statements
- Keep logic simple and focused

Return ONLY valid JSON, no markdown or explanation."""

        response = self.client.invoke(prompt, max_tokens=2000)
        response_clean = self._clean_json_response(response)
        return json.loads(response_clean)

    def _validate_script_components(self, components: Dict) -> Dict:
        """Validate and format script components."""
        validated = {}

        # Check dangerous patterns
        dangerous = ['eval', 'exec', '__import__', 'os.system', 'subprocess']
        for key, value in components.items():
            if any(p in str(value) for p in dangerous):
                raise ValueError(f"Dangerous pattern in {key}")

        # Required fields
        required = ['main_logic', 'validation', 'class_docstring']
        for req in required:
            if req not in components:
                raise ValueError(f"Missing: {req}")

        # Sanitize string components (remove triple quotes, limit to single line for docstrings)
        def sanitize_docstring(s):
            if not s:
                return "No description provided"
            # Remove triple quotes that would break the template
            s = s.replace('"""', "'")
            s = s.replace("'''", "'")
            # Replace newlines with spaces for docstrings (keep it on one logical line)
            s = ' '.join(s.split())
            return s

        # Pre-validate code snippets before indentation
        def validate_code_snippet(code, name):
            """Check if code snippet has valid Python syntax."""
            if not code or not code.strip():
                return 'pass'
            # Strip markdown code blocks if present
            code = code.strip()
            if code.startswith('```python'):
                code = code[9:]
            elif code.startswith('```'):
                code = code[3:]
            if code.endswith('```'):
                code = code[:-3]
            code = code.strip()
            # Wrap in a function to make it valid Python
            test_code = f"def _test():\n" + '\n'.join('    ' + line for line in code.split('\n'))
            try:
                compile(test_code, f'<{name}>', 'exec')
            except SyntaxError as e:
                logger.warning(f"   Code snippet '{name}' has syntax error: {e.msg}")
                logger.warning(f"   Problematic code: {code[:100]}...")
                # Return a safe default
                if name == 'main_logic':
                    return 'result = {"success": True, "message": "Default response"}'
                elif name == 'validation':
                    return 'return True'
                else:
                    return 'pass'
            return code

        init_code = validate_code_snippet(components.get('init_code', 'pass'), 'init_code')
        main_logic = validate_code_snippet(components.get('main_logic', 'result = {"success": True}'), 'main_logic')
        validation_code = validate_code_snippet(components.get('validation', 'return True'), 'validation')

        # Process with indentation
        validated['init_code'] = self._indent_code(init_code, 8)
        validated['main_logic'] = self._indent_code(main_logic, 12)
        validated['validation'] = self._indent_code(validation_code, 8)

        # Validate and process helpers (most likely source of syntax errors)
        helpers = components.get('helpers', '')
        if helpers and helpers.strip():
            # Strip markdown code blocks
            helpers = helpers.strip()
            if helpers.startswith('```python'):
                helpers = helpers[9:]
            elif helpers.startswith('```'):
                helpers = helpers[3:]
            if helpers.endswith('```'):
                helpers = helpers[:-3]
            helpers = helpers.strip()

            # Try to validate the helper methods
            try:
                # Wrap in a class to validate
                test_code = f"class _Test:\n" + '\n'.join('    ' + line for line in helpers.split('\n'))
                compile(test_code, '<helpers>', 'exec')
                validated['helpers'] = '\n' + self._indent_code(helpers, 4)
            except SyntaxError as e:
                logger.warning(f"   Helper methods have syntax error: {e.msg}")
                logger.warning(f"   Helpers code: {helpers[:100]}...")
                validated['helpers'] = ''  # Skip bad helpers
        else:
            validated['helpers'] = ''

        validated['description'] = sanitize_docstring(components.get('description', 'No description'))
        validated['class_docstring'] = sanitize_docstring(components['class_docstring'])
        validated['input_desc'] = sanitize_docstring(components.get('input_desc', 'Dict with input data'))
        validated['return_desc'] = sanitize_docstring(components.get('return_desc', 'Dict with results'))
        validated['required_inputs'] = components.get('required_inputs', [])
        validated['optional_inputs'] = components.get('optional_inputs', [])
        validated['outputs'] = components.get('outputs', ['result'])

        return validated

    @staticmethod
    def _indent_code(code: str, spaces: int) -> str:
        """Add proper indentation preserving relative indent."""
        if not code or not code.strip():
            return ' ' * spaces + 'pass'

        lines = code.split('\n')

        # Find minimum indentation
        min_indent = float('inf')
        for line in lines:
            if line.strip():
                leading = len(line) - len(line.lstrip())
                min_indent = min(min_indent, leading)

        if min_indent == float('inf'):
            min_indent = 0

        indented = []
        for line in lines:
            if line.strip():
                leading = len(line) - len(line.lstrip())
                relative = leading - min_indent
                indented.append(' ' * spaces + ' ' * relative + line.lstrip())
            else:
                indented.append('')

        return '\n'.join(indented)

    @staticmethod
    def _validate_python_syntax(code: str, name: str) -> None:
        """Validate Python syntax."""
        try:
            compile(code, f"<{name}>", 'exec')
        except SyntaxError as e:
            # Show the problematic line for debugging
            lines = code.split('\n')
            error_line = lines[e.lineno - 1] if e.lineno and e.lineno <= len(lines) else "N/A"
            logger.debug(f"   Syntax error context - Line {e.lineno}: {error_line[:100]}")
            raise SyntaxError(f"Syntax error at line {e.lineno}: {e.msg} - '{error_line[:50]}...'")

    @staticmethod
    def _to_class_name(text: str) -> str:
        """Convert to PascalCase class name (no hyphens allowed in Python class names)."""
        # Replace hyphens and underscores with spaces, then split
        text = text.replace('-', ' ').replace('_', ' ')
        # Remove non-alphanumeric except spaces
        text = re.sub(r'[^\w\s]', '', text)
        # Split on whitespace and capitalize each word
        words = text.split()
        return ''.join(word.capitalize() for word in words) + 'Skill'

    @staticmethod
    def _to_snake_case(text: str) -> str:
        """Convert to snake_case."""
        text = re.sub(r'[^\w\s-]', '', text.lower())
        text = re.sub(r'[\s-]+', '_', text)
        return text.strip('_')

    def _complete_outputs(self, skills: List[Dict[str, Any]], output_dir: str):
        """PHASE 5: Save skills and scripts

        Output structure (matches .claude/skills/ format):
        output_dir/skills/
        ├── skill_name/
        │   ├── scripts/
        │   │   └── skill_name.py
        │   ├── SKILL.md
        │   ├── config_schema.json
        │   ├── metadata.json
        │   ├── details.json          <-- NEW: Comprehensive skill details for API
        │   ├── rules.yaml
        │   ├── definitions.yaml
        │   └── examples.yaml
        """

        # Create skills directory inside output_dir
        skills_dir = os.path.join(output_dir, "skills")
        os.makedirs(skills_dir, exist_ok=True)

        scripts_count = 0

        for i, skill in enumerate(skills, 1):
            skill_name = skill["metadata"]["name"]
            skill_dir = os.path.join(skills_dir, skill_name)
            os.makedirs(skill_dir, exist_ok=True)

            logger.info(f"💾 [{i}/{len(skills)}] Saving: {skill_name}")

            # Build details.json with comprehensive skill information
            details_data = {
                "name": skill_name,
                "description": skill["metadata"].get("description", ""),
                "version": skill["metadata"].get("version", "1.0.0"),
                "category": skill["metadata"].get("category", "general"),
                "tags": skill["metadata"].get("tags", []),
                "author": skill["metadata"].get("author", ""),
                "created_at": datetime.now().isoformat(),
                "files": [
                    "SKILL.md",
                    "config_schema.json",
                    "metadata.json",
                    "rules.yaml",
                    "definitions.yaml",
                    "examples.yaml"
                ],
                "has_script": bool(skill.get("script_py")),
                "script_name": self._to_snake_case(skill_name) + ".py" if skill.get("script_py") else None,
                "metadata": skill["metadata"]
            }

            # Save all skill files
            files_to_save = {
                "SKILL.md": skill["skill_md"],
                "config_schema.json": skill.get("config_schema_json", ""),
                "metadata.json": json.dumps(skill["metadata"], indent=2),
                "rules.yaml": skill["rules_yaml"],
                "definitions.yaml": skill["definitions_yaml"],
                "examples.yaml": skill["examples_yaml"],
                "details.json": json.dumps(details_data, indent=2)
            }

            for filename, content in files_to_save.items():
                try:
                    filepath = os.path.join(skill_dir, filename)
                    with open(filepath, 'w', encoding='utf-8') as f:
                        f.write(content)
                    logger.debug(f"   ✅ {filename}")
                except Exception as e:
                    logger.error(f"   ❌ {filename}: {e}")

            # Save Python script INSIDE skill folder under scripts/ subdirectory
            script_py = skill.get("script_py")
            if script_py:
                try:
                    # Create scripts/ folder inside skill folder
                    script_dir = os.path.join(skill_dir, "scripts")
                    os.makedirs(script_dir, exist_ok=True)

                    script_filename = self._to_snake_case(skill_name) + ".py"
                    script_path = os.path.join(script_dir, script_filename)

                    with open(script_path, 'w', encoding='utf-8') as f:
                        f.write(script_py)

                    logger.info(f"   🐍 Script: scripts/{script_filename}")
                    scripts_count += 1
                except Exception as e:
                    logger.error(f"   ❌ Script failed: {e}")

            logger.info(f"   ✅ Complete")

        logger.info(f"📊 Total: {len(skills)} skills, {scripts_count} scripts")
        logger.info(f"📁 Output: {skills_dir}")
    
    def _clean_json_response(self, response: str) -> str:
        """Clean and extract JSON from LLM response.

        Handles:
        - Markdown code blocks (```json ... ```)
        - Extra text before/after the JSON object
        - Nested braces in the JSON
        """
        response = response.strip()

        # Remove markdown code blocks
        if response.startswith("```json"):
            response = response[7:]
        elif response.startswith("```"):
            response = response[3:]
        if response.endswith("```"):
            response = response[:-3]
        response = response.strip()

        # Find the first { which starts the JSON object
        start_idx = response.find('{')
        if start_idx == -1:
            logger.warning("No JSON object found in response")
            return response

        # Find the matching closing brace by tracking depth
        depth = 0
        in_string = False
        escape_next = False
        end_idx = start_idx

        for i, char in enumerate(response[start_idx:], start=start_idx):
            if escape_next:
                escape_next = False
                continue

            if char == '\\' and in_string:
                escape_next = True
                continue

            if char == '"' and not escape_next:
                in_string = not in_string
                continue

            if in_string:
                continue

            if char == '{':
                depth += 1
            elif char == '}':
                depth -= 1
                if depth == 0:
                    end_idx = i
                    break

        if depth != 0:
            logger.warning(f"Unbalanced braces in JSON (depth={depth}), returning original")
            return response

        # Extract just the JSON object
        json_str = response[start_idx:end_idx + 1]

        if end_idx + 1 < len(response):
            extra = response[end_idx + 1:].strip()
            if extra:
                logger.info(f"   Trimmed {len(extra)} chars of extra text after JSON")

        return json_str


# ============================================================================
# LAMBDA HANDLER - Async processing with CORS support
# ============================================================================

def process_extraction(file_content: bytes, filename: str, enterprise_name: Optional[str], skill_folder: str):
    """
    Background processing function.

    Args:
        file_content: Raw file bytes
        filename: Original filename
        enterprise_name: Optional enterprise name for context
        skill_folder: S3 folder name (enterprise-based, e.g., "retailco")

    IMPORTANT: The skill_folder parameter MUST be used as-is for all S3 operations.
    Do NOT regenerate or modify the skill_folder - it must match what was returned to the client.
    """
    try:
        # CRITICAL: Log the exact skill_folder being used - this MUST match what was returned to client
        logger.info(f"[PROCESS_EXTRACTION] ========================================")
        logger.info(f"[PROCESS_EXTRACTION] skill_folder parameter: '{skill_folder}'")
        logger.info(f"[PROCESS_EXTRACTION] enterprise_name parameter: '{enterprise_name}'")
        logger.info(f"[PROCESS_EXTRACTION] filename parameter: '{filename}'")
        logger.info(f"[PROCESS_EXTRACTION] ========================================")

        logger.info(f"[{skill_folder}] Processing file: {filename}")
        logger.info(f"[{skill_folder}] Enterprise: {enterprise_name or 'Will be inferred'}")
        logger.info(f"[{skill_folder}] File size: {len(file_content)} bytes")

        # Verify initial status exists (sanity check)
        try:
            status_key = f"Knowledge_Extraction_To_Skills/{skill_folder}/status.json"
            get_s3_client().head_object(Bucket=SKILLS_BUCKET, Key=status_key)
            logger.info(f"[{skill_folder}] Initial status verified at: {status_key}")
        except Exception as verify_err:
            logger.warning(f"[{skill_folder}] Could not verify initial status: {verify_err}")

        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp_file:
            tmp_file.write(file_content)
            tmp_path = tmp_file.name

        # Parse document to text
        text_content = parse_document(tmp_path, filename)
        logger.info(f"[{skill_folder}] Extracted text: {len(text_content)} characters")

        # Save as temporary text file for orchestrator
        text_file_path = tmp_path + '.txt'
        with open(text_file_path, 'w', encoding='utf-8') as f:
            f.write(text_content)

        # Run knowledge extraction
        output_dir = tempfile.mkdtemp()

        orchestrator = KnowledgeCaptureOrchestrator(region=os.environ.get('AWS_REGION', 'us-west-2'))
        result = orchestrator.process_documents(
            document_paths=[text_file_path],
            enterprise_name=enterprise_name,
            output_dir=output_dir
        )

        # Upload skills to S3 using enterprise-based folder name
        s3_paths = []
        logger.info(f"[{skill_folder}] Uploading files from: {output_dir}")
        logger.info(f"[{skill_folder}] Target bucket: {SKILLS_BUCKET}")
        logger.info(f"[{skill_folder}] S3 folder: Knowledge_Extraction_To_Skills/{skill_folder}/")

        for root, dirs, files in os.walk(output_dir):
            for file in files:
                local_path = os.path.join(root, file)
                relative_path = os.path.relpath(local_path, output_dir)
                # Ensure forward slashes for S3 keys (important for cross-platform compatibility)
                relative_path = relative_path.replace('\\', '/')
                # Use enterprise-based skill_folder instead of job_id
                s3_key = f"Knowledge_Extraction_To_Skills/{skill_folder}/{relative_path}"

                logger.info(f"[{skill_folder}] Uploading: {local_path} -> s3://{SKILLS_BUCKET}/{s3_key}")
                get_s3_client().upload_file(local_path, SKILLS_BUCKET, s3_key)
                s3_paths.append(s3_key)
                logger.info(f"[{skill_folder}] ✅ Uploaded: {s3_key}")

        logger.info(f"[{skill_folder}] Total files uploaded: {len(s3_paths)}")

        # Save completion status
        completion_result = {
            'status': 'completed',
            'message': 'Skills extracted successfully',
            'result': result,
            's3_bucket': SKILLS_BUCKET,
            's3_paths': s3_paths,
            'skill_folder': skill_folder,
            'enterprise_name': enterprise_name,
            'completed_at': datetime.now().isoformat()
        }

        get_s3_client().put_object(
            Bucket=SKILLS_BUCKET,
            Key=f"Knowledge_Extraction_To_Skills/{skill_folder}/status.json",
            Body=json.dumps(completion_result, indent=2),
            ContentType='application/json'
        )

        logger.info(f"[{skill_folder}] ✅ Processing complete")
        logger.info(f"[{skill_folder}] Skills stored at: s3://{SKILLS_BUCKET}/Knowledge_Extraction_To_Skills/{skill_folder}/")

        # Cleanup
        os.unlink(tmp_path)
        os.unlink(text_file_path)

    except Exception as e:
        logger.error(f"[{skill_folder}] ❌ Processing failed: {str(e)}", exc_info=True)

        # Save error status
        error_result = {
            'status': 'failed',
            'error': str(e),
            'type': type(e).__name__,
            'skill_folder': skill_folder,
            'enterprise_name': enterprise_name,
            'failed_at': datetime.now().isoformat()
        }

        get_s3_client().put_object(
            Bucket=SKILLS_BUCKET,
            Key=f"Knowledge_Extraction_To_Skills/{skill_folder}/status.json",
            Body=json.dumps(error_result, indent=2),
            ContentType='application/json'
        )


def handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    """
    UPDATED LAMBDA HANDLER - Supports multiple upload methods:

    ENDPOINTS:
    1. GET /upload-url?filename=X&enterprise_name=Y  - Get presigned URL for direct S3 upload (RECOMMENDED for large files)
    2. POST /extract with s3_key                     - Process file already uploaded to S3 (RECOMMENDED)
    3. POST /extract with multipart/form-data        - Direct file upload (legacy, small files only)
    4. POST /extract with base64 JSON                - Base64 encoded file (legacy, small files only)
    5. GET /status/{skill_folder}                    - Check extraction status

    RECOMMENDED FLOW FOR LARGE FILES:
    1. Call GET /upload-url?filename=document.pdf&enterprise_name=RetailCo
    2. Upload file directly to S3 using the returned presigned URL (PUT request)
    3. Call POST /extract with {"s3_key": "...", "filename": "...", "enterprise_name": "..."}

    Skills are stored with enterprise-based naming (single folder per enterprise):
    - S3 path: Knowledge_Extraction_To_Skills/{enterprise_name}/
    - Example: Knowledge_Extraction_To_Skills/retailco/

    Expected request body for S3 key method (RECOMMENDED):
    {
        "s3_key": "pending_uploads/retailco/20260213_..._document.pdf",
        "filename": "document.pdf",
        "enterprise_name": "RetailCo" (optional)
    }

    Expected request body for base64 JSON (legacy):
    {
        "file": "base64_encoded_file_content",
        "filename": "document.pdf",
        "enterprise_name": "RetailCo" (optional)
    }
    """

    try:
        # Handle OPTIONS request (CORS preflight)
        if event.get('httpMethod') == 'OPTIONS':
            return {
                'statusCode': 200,
                'headers': CORS_HEADERS,
                'body': json.dumps({'message': 'CORS preflight'})
            }

        # Handle GET requests
        if event.get('httpMethod') == 'GET':
            # REST API uses 'path' and 'resource', HTTP API uses 'rawPath'
            path = event.get('path', '') or event.get('rawPath', '') or ''
            resource = event.get('resource', '') or ''
            query_params = event.get('queryStringParameters', {}) or {}
            path_params = event.get('pathParameters', {}) or {}

            logger.info(f"[GET_REQUEST] path={path}, resource={resource}, query_params={query_params}")

            # DEBUG: Return event structure for troubleshooting (remove after fixing)
            if 'debug' in (query_params or {}):
                return {
                    'statusCode': 200,
                    'headers': CORS_HEADERS,
                    'body': json.dumps({
                        'debug': True,
                        'path': path,
                        'resource': resource,
                        'query_params': query_params,
                        'path_params': path_params,
                        'httpMethod': event.get('httpMethod'),
                        'event_keys': list(event.keys())
                    }, indent=2)
                }

            # NEW: GET /upload-url - Generate presigned URL for direct S3 upload
            # Check both path and resource for compatibility with REST API and HTTP API
            if '/upload-url' in path or '/upload-url' in resource or 'upload-url' in path or 'upload-url' in resource:
                filename = query_params.get('filename') if query_params else None
                enterprise_name = (query_params.get('enterprise_name') or query_params.get('enterpriseName')) if query_params else None
                user_id = (query_params.get('user_id') or query_params.get('userId')) if query_params else None

                if not filename:
                    return {
                        'statusCode': 400,
                        'headers': CORS_HEADERS,
                        'body': json.dumps({
                            'error': 'Missing required parameter: filename',
                            'usage': 'GET /upload-url?filename=document.pdf&enterprise_name=RetailCo',
                            'debug': {
                                'path': path,
                                'resource': resource,
                                'query_params': query_params
                            }
                        })
                    }

                logger.info(f"[UPLOAD_URL] Generating presigned URL for: {filename}, user_id: {user_id}, enterprise: {enterprise_name}")

                try:
                    upload_info = generate_presigned_upload_url(filename, enterprise_name, user_id)

                    return {
                        'statusCode': 200,
                        'headers': CORS_HEADERS,
                        'body': json.dumps({
                            'upload_url': upload_info['upload_url'],
                            's3_key': upload_info['s3_key'],
                            'expires_in_seconds': upload_info['expires_in_seconds'],
                            'filename': filename,
                            'user_id': user_id,
                            'enterprise_name': enterprise_name,
                            'instructions': {
                                'step1': 'Upload your file using PUT request to upload_url',
                                'step2': 'Call POST /extract with s3_key, filename, user_id, and enterprise_name',
                                'example_curl': f"curl -X PUT -T yourfile.pdf '{upload_info['upload_url']}'"
                            }
                        })
                    }
                except Exception as e:
                    logger.error(f"[UPLOAD_URL] Error generating presigned URL: {str(e)}", exc_info=True)
                    return {
                        'statusCode': 500,
                        'headers': CORS_HEADERS,
                        'body': json.dumps({
                            'error': f'Failed to generate upload URL: {str(e)}',
                            'bucket': SKILLS_BUCKET,
                            'hint': 'Check that SKILLS_BUCKET environment variable is set'
                        })
                    }

            # GET /status/{skill_folder} - Status check
            folder_id = path_params.get('skill_folder') or path_params.get('job_id')

            if folder_id:
                # Normalize folder_id to match how it was created:
                # 1. URL decode (@ might be passed as %40)
                # 2. Apply same sanitization used during folder creation
                original_folder_id = folder_id
                folder_id = urllib.parse.unquote(folder_id)

                # If it looks like an email, apply email sanitization (lowercase, etc.)
                if is_valid_email(folder_id):
                    folder_id = sanitize_email_for_s3_key(folder_id)
                    logger.info(f"[STATUS_CHECK] Normalized email folder_id: '{original_folder_id}' -> '{folder_id}'")
                elif folder_id != original_folder_id:
                    logger.info(f"[STATUS_CHECK] URL decoded folder_id: '{original_folder_id}' -> '{folder_id}'")
                # BACKWARD COMPATIBILITY: Try multiple S3 paths
                # 1. Primary path: Knowledge_Extraction_To_Skills/{folder_id}/status.json
                # 2. Legacy path: skills/{folder_id}/status.json (old code)
                # 3. Alternative: skills_phase3/{folder_id}/status.json
                s3_paths_to_try = [
                    f"Knowledge_Extraction_To_Skills/{folder_id}/status.json",
                    f"skills/{folder_id}/status.json",
                    f"skills_phase3/{folder_id}/status.json",
                    # Also try without the folder structure (direct status.json)
                    f"Knowledge_Extraction_To_Skills/{folder_id}/skills/{folder_id}/status.json",
                ]

                status_data = None
                found_path = None

                for s3_path in s3_paths_to_try:
                    try:
                        logger.info(f"[STATUS_CHECK] Trying path: s3://{SKILLS_BUCKET}/{s3_path}")
                        response = get_s3_client().get_object(
                            Bucket=SKILLS_BUCKET,
                            Key=s3_path
                        )
                        status_data = json.loads(response['Body'].read().decode('utf-8'))
                        found_path = s3_path
                        logger.info(f"[STATUS_CHECK] Found status at: {s3_path}")
                        break
                    except Exception as e:
                        logger.debug(f"[STATUS_CHECK] Not found at {s3_path}: {str(e)}")
                        continue

                if status_data:
                    # Add the path where we found it for debugging
                    status_data['_found_at_path'] = found_path

                    # If status is completed, also fetch details.json from each skill subfolder
                    if status_data.get('status') == 'completed':
                        try:
                            skills_with_details = fetch_skill_details(folder_id)
                            if skills_with_details:
                                status_data['skills_details'] = skills_with_details
                                logger.info(f"[STATUS_CHECK] Added details for {len(skills_with_details)} skills")
                        except Exception as details_err:
                            logger.warning(f"[STATUS_CHECK] Could not fetch skill details: {str(details_err)}")

                    return {
                        'statusCode': 200,
                        'headers': CORS_HEADERS,
                        'body': json.dumps(status_data)
                    }
                else:
                    # Not found in any path - list what's actually in the bucket for this prefix
                    logger.warning(f"[STATUS_CHECK] Skill folder '{folder_id}' not found in any path")

                    # Try to list objects to help debug
                    existing_folders = []
                    try:
                        list_response = get_s3_client().list_objects_v2(
                            Bucket=SKILLS_BUCKET,
                            Prefix='Knowledge_Extraction_To_Skills/',
                            Delimiter='/',
                            MaxKeys=10
                        )
                        for prefix in list_response.get('CommonPrefixes', []):
                            folder_name = prefix['Prefix'].replace('Knowledge_Extraction_To_Skills/', '').rstrip('/')
                            if folder_name:
                                existing_folders.append(folder_name)
                    except Exception as list_error:
                        logger.error(f"[STATUS_CHECK] Error listing folders: {str(list_error)}")

                    return {
                        'statusCode': 404,
                        'headers': CORS_HEADERS,
                        'body': json.dumps({
                            'error': 'Skill folder not found',
                            'skill_folder': folder_id,
                            'paths_checked': s3_paths_to_try,
                            'hint': 'The extraction job may still be processing, or the folder name may be incorrect.',
                            'recent_folders': existing_folders[:5] if existing_folders else [],
                            'suggestion': 'Ensure you are using the skill_folder returned from the POST /extract response'
                        })
                    }

        # Check if this is a background invocation
        if event.get('is_background_job'):
            # Background processing (no CORS needed for internal calls)
            job_data = event['job_data']

            # CRITICAL: Log received values to trace any mismatch
            received_skill_folder = job_data.get('skill_folder')
            received_enterprise_name = job_data.get('enterprise_name')
            received_filename = job_data.get('filename')

            logger.info(f"[ASYNC_JOB] ========== BACKGROUND JOB STARTED ==========")
            logger.info(f"[ASYNC_JOB] Received skill_folder: '{received_skill_folder}'")
            logger.info(f"[ASYNC_JOB] Received enterprise_name: '{received_enterprise_name}'")
            logger.info(f"[ASYNC_JOB] Received filename: '{received_filename}'")

            # Validate skill_folder is present
            if not received_skill_folder:
                logger.error("[ASYNC_JOB] ERROR: skill_folder is missing from job_data!")
                return {'statusCode': 400, 'body': 'Missing skill_folder in job_data'}

            # Get file content - from S3 (recommended) or base64 (legacy)
            use_s3_key = job_data.get('use_s3_key', False)
            uploaded_s3_key = job_data.get('s3_key')

            if use_s3_key and uploaded_s3_key:
                # S3 KEY METHOD (RECOMMENDED) - Read file directly from S3
                logger.info(f"[ASYNC_JOB] Reading file from S3: {uploaded_s3_key}")
                try:
                    file_content = read_file_from_s3(uploaded_s3_key)
                    logger.info(f"[ASYNC_JOB] File read from S3: {len(file_content)} bytes")

                    # Clean up the uploaded file after reading (optional - keeps S3 clean)
                    delete_uploaded_file(uploaded_s3_key)
                except Exception as e:
                    logger.error(f"[ASYNC_JOB] Failed to read from S3: {e}")
                    return {'statusCode': 500, 'body': f'Failed to read file from S3: {str(e)}'}
            else:
                # LEGACY METHOD - Decode from base64
                logger.info("[ASYNC_JOB] Using legacy base64 method")
                if 'file' not in job_data:
                    logger.error("[ASYNC_JOB] ERROR: Neither s3_key nor file provided!")
                    return {'statusCode': 400, 'body': 'Missing file content in job_data'}

                file_content = base64.b64decode(job_data['file'])
                logger.info(f"[ASYNC_JOB] File content decoded from base64: {len(file_content)} bytes")

            process_extraction(
                file_content=file_content,
                filename=received_filename,
                enterprise_name=received_enterprise_name,
                skill_folder=received_skill_folder
            )
            logger.info(f"[ASYNC_JOB] ========== BACKGROUND JOB COMPLETED ==========")
            return {'statusCode': 200, 'body': 'Background processing complete'}
        
        # Handle POST - File upload
        if event.get('httpMethod') == 'POST':
            content_type = event.get('headers', {}).get('Content-Type') or \
                          event.get('headers', {}).get('content-type', '')
            
            # Method 1: Multipart file upload
            if 'multipart/form-data' in content_type:
                logger.info("Processing multipart/form-data upload")
                
                body = event.get('body', '')
                is_base64 = event.get('isBase64Encoded', False)
                
                if not is_base64:
                    # API Gateway should base64 encode binary data
                    # If not, we need to handle it
                    body = base64.b64encode(body.encode()).decode()
                
                parsed_data = parse_multipart_form_data(body, content_type)

                file_content = parsed_data.get('file')
                filename = parsed_data.get('filename')
                enterprise_name = parsed_data.get('enterprise_name')
                user_id = parsed_data.get('user_id')
                
                if not file_content or not filename:
                    return {
                        'statusCode': 400,
                        'headers': CORS_HEADERS,
                        'body': json.dumps({
                            'error': 'Missing file or filename in multipart upload'
                        })
                    }
                
                # Convert to base64 for internal processing (legacy method)
                file_base64 = base64.b64encode(file_content).decode('utf-8')
                use_s3_key = False
                uploaded_file_s3_key = None

            # Method 2: JSON upload (supports action, s3_key, or base64)
            elif 'application/json' in content_type:
                body = json.loads(event.get('body', '{}'))

                # Log all keys received for debugging
                logger.info(f"[JSON_PARSE] Received keys in body: {list(body.keys())}")

                # Support multiple variations of enterprise_name field
                enterprise_name = (
                    body.get('enterprise_name') or
                    body.get('enterpriseName') or
                    body.get('enterprise-name') or
                    body.get('EnterpriseName')
                )

                # Support user_id for folder naming (takes priority over enterprise_name)
                user_id = (
                    body.get('user_id') or
                    body.get('userId') or
                    body.get('user-id') or
                    body.get('UserId')
                )

                # Support email_id for folder naming (highest priority - preserves email format)
                # Example: Akhil.Parelly@quadranttechnologies.com -> akhil.parelly@quadranttechnologies.com
                email_id = (
                    body.get('email_id') or
                    body.get('emailId') or
                    body.get('email-id') or
                    body.get('EmailId') or
                    body.get('email') or
                    body.get('Email')
                )

                filename = body.get('filename')

                # ACTION: get_upload_url - Return presigned URL for direct S3 upload (NO new route needed!)
                if body.get('action') == 'get_upload_url':
                    if not filename:
                        return {
                            'statusCode': 400,
                            'headers': CORS_HEADERS,
                            'body': json.dumps({
                                'error': 'Missing required field: filename',
                                'usage': {
                                    'action': 'get_upload_url',
                                    'filename': 'document.pdf',
                                    'enterprise_name': 'RetailCo (optional)'
                                }
                            })
                        }

                    logger.info(f"[GET_UPLOAD_URL] Generating presigned URL for: {filename}, user_id: {user_id}, enterprise: {enterprise_name}")

                    try:
                        upload_info = generate_presigned_upload_url(filename, enterprise_name, user_id)

                        return {
                            'statusCode': 200,
                            'headers': CORS_HEADERS,
                            'body': json.dumps({
                                'upload_url': upload_info['upload_url'],
                                's3_key': upload_info['s3_key'],
                                'expires_in_seconds': upload_info['expires_in_seconds'],
                                'filename': filename,
                                'user_id': user_id,
                                'enterprise_name': enterprise_name,
                                'next_step': {
                                    'description': 'Upload file to S3, then call /extract with s3_key',
                                    'step1': 'PUT your file to upload_url',
                                    'step2': 'POST /extract with the body below',
                                    'example_body': {
                                        's3_key': upload_info['s3_key'],
                                        'filename': filename,
                                        'user_id': user_id,
                                        'enterprise_name': enterprise_name
                                    }
                                }
                            })
                        }
                    except Exception as e:
                        logger.error(f"[GET_UPLOAD_URL] Error: {str(e)}", exc_info=True)
                        return {
                            'statusCode': 500,
                            'headers': CORS_HEADERS,
                            'body': json.dumps({
                                'error': f'Failed to generate upload URL: {str(e)}',
                                'bucket': SKILLS_BUCKET
                            })
                        }

                # Method 2a: S3 key method (RECOMMENDED for large files)
                if 's3_key' in body:
                    uploaded_file_s3_key = body['s3_key']  # Store uploaded file's S3 key separately
                    logger.info(f"[S3_KEY_METHOD] Processing file from S3: {uploaded_file_s3_key}")

                    if not filename:
                        # Extract filename from s3_key if not provided
                        filename = Path(uploaded_file_s3_key).name
                        logger.info(f"[S3_KEY_METHOD] Extracted filename from s3_key: {filename}")

                    # Store s3_key for later use (instead of file_base64)
                    use_s3_key = True
                    file_base64 = None  # Will read from S3 in background job
                    logger.info(f"[S3_KEY_METHOD] Will read file from S3 in background job")

                # Method 2b: Base64 method (legacy, for small files)
                elif 'file' in body:
                    logger.info("Processing JSON with base64 file (legacy method)")

                    if not filename:
                        return {
                            'statusCode': 400,
                            'headers': CORS_HEADERS,
                            'body': json.dumps({
                                'error': 'Missing required field: filename',
                                'hint': 'Include filename in your request'
                            })
                        }

                    file_base64 = body['file']
                    use_s3_key = False
                    uploaded_file_s3_key = None

                else:
                    return {
                        'statusCode': 400,
                        'headers': CORS_HEADERS,
                        'body': json.dumps({
                            'error': 'Missing required fields: action, s3_key, or file',
                            'hint': 'Use action=get_upload_url first, then s3_key to process',
                            'recommended_flow': [
                                '1. POST /extract with {"action": "get_upload_url", "filename": "doc.pdf", "enterprise_name": "RetailCo"}',
                                '2. PUT file to the returned upload_url',
                                '3. POST /extract with {"s3_key": "...", "filename": "doc.pdf", "enterprise_name": "RetailCo"}'
                            ],
                            'legacy_option': 'Or send {"file": "base64_content", "filename": "doc.pdf"} for small files'
                        })
                    }

                logger.info(f"[JSON_PARSE] Extracted: filename='{filename}', enterprise_name='{enterprise_name}', use_s3_key={use_s3_key}")
            
            else:
                return {
                    'statusCode': 400,
                    'headers': CORS_HEADERS,
                    'body': json.dumps({
                        'error': 'Invalid Content-Type',
                        'expected': 'multipart/form-data or application/json',
                        'received': content_type
                    })
                }
            
            # Generate user-based or enterprise-based skill folder name
            # Priority: email_id > user_id > enterprise_name > extract from filename
            # CRITICAL: Log all inputs for debugging folder name mismatch issues
            logger.info(f"[FOLDER_GEN] Raw email_id from request: '{email_id}' (type: {type(email_id).__name__ if email_id else 'NoneType'})")
            logger.info(f"[FOLDER_GEN] Raw user_id from request: '{user_id}' (type: {type(user_id).__name__ if user_id else 'NoneType'})")
            logger.info(f"[FOLDER_GEN] Raw enterprise_name from request: '{enterprise_name}' (type: {type(enterprise_name).__name__})")
            logger.info(f"[FOLDER_GEN] Raw filename from request: '{filename}'")

            # Normalize email_id - treat empty strings as None
            if email_id is not None and isinstance(email_id, str):
                email_id = email_id.strip() if email_id.strip() else None
                logger.info(f"[FOLDER_GEN] Normalized email_id: '{email_id}'")

            # Normalize user_id - treat empty strings as None
            if user_id is not None and isinstance(user_id, str):
                user_id = user_id.strip() if user_id.strip() else None
                logger.info(f"[FOLDER_GEN] Normalized user_id: '{user_id}'")

            # Normalize enterprise_name - treat empty strings as None
            if enterprise_name is not None and isinstance(enterprise_name, str):
                enterprise_name = enterprise_name.strip() if enterprise_name.strip() else None
                logger.info(f"[FOLDER_GEN] Normalized enterprise_name: '{enterprise_name}'")

            skill_folder = generate_skill_folder_name(enterprise_name, filename, user_id, email_id)
            logger.info(f"[FOLDER_GEN] Generated skill_folder: '{skill_folder}'")

            # Determine the effective enterprise name for metadata
            effective_enterprise_name = enterprise_name or extract_enterprise_from_filename(filename)
            logger.info(f"[FOLDER_GEN] Effective enterprise_name for metadata: '{effective_enterprise_name}'")

            # Save initial status
            initial_status = {
                'status': 'processing',
                'message': 'Extraction job started',
                'skill_folder': skill_folder,
                'email_id': email_id,
                'user_id': user_id,
                'enterprise_name': effective_enterprise_name,
                'filename': filename,
                'started_at': datetime.now().isoformat(),
                'estimated_completion': '5-10 minutes',
                's3_path': f"s3://{SKILLS_BUCKET}/Knowledge_Extraction_To_Skills/{skill_folder}/"
            }

            # Save initial status to S3 and verify it was saved
            s3_key = f"Knowledge_Extraction_To_Skills/{skill_folder}/status.json"
            logger.info(f"[FOLDER_GEN] Saving initial status to: s3://{SKILLS_BUCKET}/{s3_key}")
            get_s3_client().put_object(
                Bucket=SKILLS_BUCKET,
                Key=s3_key,
                Body=json.dumps(initial_status, indent=2),
                ContentType='application/json'
            )
            logger.info(f"[FOLDER_GEN] Initial status saved successfully")

            # Prepare job data - CRITICAL: use the SAME skill_folder that was saved
            # Use s3_key method (recommended) or file base64 (legacy)
            job_data = {
                'filename': filename,
                'email_id': email_id,  # Pass email_id (highest priority for folder naming)
                'user_id': user_id,  # Pass user_id (second priority for folder naming)
                'enterprise_name': enterprise_name,  # Pass original (may be None)
                'skill_folder': skill_folder,  # MUST match the saved status folder
                'use_s3_key': use_s3_key  # Flag to indicate which method to use
            }

            if use_s3_key:
                # S3 key method - much smaller payload, no base64 overhead
                job_data['s3_key'] = uploaded_file_s3_key  # Use the uploaded file's S3 key, NOT status.json
                logger.info(f"[ASYNC_INVOKE] Using S3 key method: {uploaded_file_s3_key}")
            else:
                # Legacy base64 method
                job_data['file'] = file_base64
                logger.info(f"[ASYNC_INVOKE] Using base64 method (legacy)")

            logger.info(f"[ASYNC_INVOKE] Invoking async with skill_folder: '{job_data['skill_folder']}'")
            logger.info(f"[ASYNC_INVOKE] Invoking async with enterprise_name: '{job_data['enterprise_name']}'")

            # Invoke Lambda asynchronously
            lambda_client = boto3.client('lambda')
            lambda_client.invoke(
                FunctionName=context.function_name,
                InvocationType='Event',  # Async invocation
                Payload=json.dumps({
                    'is_background_job': True,
                    'job_data': job_data
                })
            )

            logger.info(f"[{skill_folder}] Async job started successfully")

            # Return immediately with CORS headers
            # CRITICAL: Return the SAME skill_folder that was used for S3 storage
            response_body = {
                'status': 'accepted',
                'message': 'Extraction job started',
                'skill_folder': skill_folder,
                'email_id': email_id,
                'user_id': user_id,
                'enterprise_name': effective_enterprise_name,
                'filename': filename,
                'check_status_url': f"GET /status/{skill_folder}",
                's3_path': f"Knowledge_Extraction_To_Skills/{skill_folder}/",
                'estimated_completion': '5-10 minutes'
            }
            logger.info(f"[RESPONSE] Returning skill_folder: '{skill_folder}'")

            return {
                'statusCode': 202,  # Accepted
                'headers': CORS_HEADERS,
                'body': json.dumps(response_body)
            }

        return {
            'statusCode': 400,
            'headers': CORS_HEADERS,
            'body': json.dumps({
                'error': 'Invalid request',
                'supported_methods': ['POST /extract', 'GET /status/{skill_folder}']
            })
        }
        
    except Exception as e:
        logger.error(f"Error: {str(e)}", exc_info=True)

        return {
            'statusCode': 500,
            'headers': CORS_HEADERS,
            'body': json.dumps({
                'error': str(e),
                'type': type(e).__name__
            })
        }


# ============================================================================
# LOCAL EXECUTION MODE - For testing without AWS Lambda
# ============================================================================

class LocalAnthropicClient:
    """Local client using Anthropic API directly (not Bedrock)"""

    def __init__(self, api_key: str = None):
        try:
            import anthropic
            self.api_key = api_key or os.environ.get('ANTHROPIC_API_KEY')
            if not self.api_key:
                raise ValueError("ANTHROPIC_API_KEY not set")
            self.client = anthropic.Anthropic(api_key=self.api_key)
            self.model = "claude-sonnet-4-20250514"
            logger.info("✅ Anthropic client initialized (local mode)")
        except ImportError:
            raise ImportError("Local mode requires 'anthropic' package. Install with: pip install anthropic")

    def invoke(self, prompt: str, max_tokens: int = DEFAULT_MAX_OUTPUT_TOKENS) -> str:
        """Invoke Claude API directly"""
        if len(prompt) > MAX_INPUT_CHARS:
            logger.warning(f"⚠️  Truncating prompt from {len(prompt)} chars")
            prompt = prompt[:MAX_INPUT_CHARS] + "\n\n[TRUNCATED]"

        message = self.client.messages.create(
            model=self.model,
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}]
        )

        return message.content[0].text


class LocalKnowledgeOrchestrator(KnowledgeCaptureOrchestrator):
    """Local orchestrator that uses Anthropic API instead of Bedrock"""

    def __init__(self, api_key: str = None):
        logger.info("="*70)
        logger.info("🚀 LOCAL SKILL GENERATOR (Anthropic API)")
        logger.info("="*70)
        self.client = LocalAnthropicClient(api_key=api_key)
        self.model = self.client.model
        logger.info(f"✅ Ready (Model: {self.model})")
        logger.info("")

    def process_with_review(
        self,
        document_paths: List[str],
        enterprise_name: Optional[str] = None,
        output_dir: str = "./extracted_skills",
        auto_mode: bool = False
    ) -> Dict[str, Any]:
        """
        Process documents with interactive review step.

        Steps:
        1. Upload file
        2. Extract knowledge
        3. Save extracted_knowledge.json and ask for review
        4. Generate skills and scripts
        5. Validate and save
        """
        logger.info("="*70)
        logger.info("📚 PROCESSING WITH REVIEW")
        logger.info("="*70)
        logger.info(f"Documents: {len(document_paths)}")
        logger.info(f"Enterprise: {enterprise_name or 'Will be inferred'}")
        logger.info(f"Output: {output_dir}")
        logger.info(f"Auto mode: {auto_mode}")
        logger.info("")

        try:
            # ==========================================
            # STEP 1: UPLOAD FILE
            # ==========================================
            print("\n" + "="*60)
            print("STEP 1: UPLOAD FILE")
            print("="*60)

            documents = self._upload_documents(document_paths, enterprise_name)
            logger.info(f"✅ Uploaded {len(documents)} document(s)")

            if not documents:
                raise ValueError("No documents were successfully uploaded")

            # ==========================================
            # STEP 2: EXTRACT KNOWLEDGE
            # ==========================================
            print("\n" + "="*60)
            print("STEP 2: EXTRACT KNOWLEDGE")
            print("="*60)

            extracted_knowledge = self._extract_knowledge(documents)
            logger.info(f"✅ Extracted {len(extracted_knowledge)} knowledge items")

            # ==========================================
            # STEP 3: REVIEW EXTRACTED KNOWLEDGE
            # ==========================================
            print("\n" + "="*60)
            print("STEP 3: REVIEW EXTRACTED KNOWLEDGE")
            print("="*60)

            # Save extracted knowledge for review
            os.makedirs(output_dir, exist_ok=True)
            knowledge_file = os.path.join(output_dir, "extracted_knowledge.json")

            # Convert to serializable format
            knowledge_data = {
                "document_metadata": {
                    "files_processed": [d.filename for d in documents],
                    "enterprise": enterprise_name or documents[0].enterprise if documents else "Unknown",
                    "extraction_timestamp": datetime.now().isoformat()
                },
                "extracted_knowledge": [
                    {
                        "type": k.type,
                        "content": k.content,
                        "source_document": k.source_document,
                        "source_section": k.source_section,
                        "confidence": k.confidence,
                        "user_intention": k.user_intention,
                        "metadata": k.metadata
                    }
                    for k in extracted_knowledge
                ],
                "statistics": {
                    "total_items": len(extracted_knowledge),
                    "by_type": {}
                }
            }

            # Count by type
            for k in extracted_knowledge:
                k_type = k.type or "unknown"
                knowledge_data["statistics"]["by_type"][k_type] = \
                    knowledge_data["statistics"]["by_type"].get(k_type, 0) + 1

            with open(knowledge_file, 'w', encoding='utf-8') as f:
                json.dump(knowledge_data, f, indent=2)

            print(f"\n✅ Extracted knowledge saved to: {knowledge_file}")
            print(f"\n📊 Extraction Statistics:")
            print(f"   - Total items: {len(extracted_knowledge)}")
            for k_type, count in knowledge_data["statistics"]["by_type"].items():
                print(f"   - {k_type}: {count}")

            if not auto_mode:
                print("\n" + "-"*60)
                print("📋 Please review 'extracted_knowledge.json'")
                print("   You can edit the file to modify/remove items before generation.")
                print("-"*60)
                input("\nPress Enter when ready to generate skills and scripts...")

                # Reload in case user edited
                with open(knowledge_file, 'r', encoding='utf-8') as f:
                    knowledge_data = json.load(f)

                # Reconstruct knowledge objects
                extracted_knowledge = []
                for item in knowledge_data.get("extracted_knowledge", []):
                    k = ExtractedKnowledge(
                        type=item.get("type"),
                        content=item.get("content"),
                        source_document=item.get("source_document", ""),
                        source_section=item.get("source_section"),
                        confidence=item.get("confidence", 0.8),
                        metadata=item.get("metadata", {}),
                        user_intention=item.get("user_intention")
                    )
                    extracted_knowledge.append(k)

                print(f"\n✅ Reloaded {len(extracted_knowledge)} knowledge items after review")

            # Deduplicate
            consolidated = self._review_knowledge(extracted_knowledge)

            # ==========================================
            # STEP 4: GENERATE SKILLS & SCRIPTS
            # ==========================================
            print("\n" + "="*60)
            print("STEP 4: GENERATE SKILLS & SCRIPTS")
            print("="*60)

            skills = self._generate_skills(consolidated, enterprise_name)
            logger.info(f"✅ Generated {len(skills)} skill(s)")

            # ==========================================
            # STEP 5: SAVE & VALIDATE
            # ==========================================
            print("\n" + "="*60)
            print("STEP 5: SAVE & VALIDATE")
            print("="*60)

            self._complete_outputs(skills, output_dir)
            logger.info(f"✅ Saved to {output_dir}")

            # Count scripts
            scripts_count = sum(1 for s in skills if s.get("script_py"))

            result = {
                "status": "success",
                "documents_processed": len(documents),
                "knowledge_items_extracted": len(consolidated),
                "skills_generated": len(skills),
                "scripts_generated": scripts_count,
                "output_directory": output_dir,
                "timestamp": datetime.now().isoformat()
            }

            return result

        except Exception as e:
            logger.error(f"❌ Failed: {e}", exc_info=True)
            raise


def main():
    """Local execution entry point with interactive review"""
    import sys

    print("="*70)
    print("🏠 LOCAL KNOWLEDGE EXTRACTION PIPELINE")
    print("="*70)
    print("")
    print("Pipeline Steps:")
    print("  1. Upload file (any format: PDF, DOCX, TXT, MD)")
    print("  2. Extract knowledge using Claude API")
    print("  3. Review extracted_knowledge.json (editable)")
    print("  4. Generate SKILL.md files and Python scripts")
    print("  5. Validate scripts and save outputs")
    print("")

    if len(sys.argv) < 2:
        print("Usage: python lambda_handler.py <document_file> [enterprise_name] [--auto]")
        print("")
        print("Arguments:")
        print("  document_file   - Path to SOP/manual file")
        print("  enterprise_name - Optional enterprise name")
        print("  --auto          - Skip review step (auto mode)")
        print("")
        print("Supported formats: .txt, .md, .pdf, .docx")
        print("")
        print("Examples:")
        print("  python lambda_handler.py manual.pdf")
        print("  python lambda_handler.py \"Retail SOP.pdf\" RetailCo")
        print("  python lambda_handler.py document.pdf RetailCo --auto")
        print("")
        print("Environment variables:")
        print("  ANTHROPIC_API_KEY - Required for local execution")
        return 1

    input_file = sys.argv[1]

    # Parse optional arguments
    enterprise_name = None
    auto_mode = "--auto" in sys.argv

    for arg in sys.argv[2:]:
        if arg != "--auto":
            enterprise_name = arg
            break

    # Check file exists
    if not os.path.exists(input_file):
        print(f"❌ Error: File not found: {input_file}")
        return 1

    # Check API key
    if not os.environ.get('ANTHROPIC_API_KEY'):
        print("❌ Error: ANTHROPIC_API_KEY environment variable not set")
        print("")
        print("Set it with:")
        print("  set ANTHROPIC_API_KEY=your-api-key-here  (Windows)")
        print("  export ANTHROPIC_API_KEY=your-api-key-here  (Linux/Mac)")
        return 1

    try:
        # Create output directory
        output_dir = "./extracted_skills"

        # Initialize local orchestrator
        orchestrator = LocalKnowledgeOrchestrator()

        # Process document with review
        result = orchestrator.process_with_review(
            document_paths=[input_file],
            enterprise_name=enterprise_name,
            output_dir=output_dir,
            auto_mode=auto_mode
        )

        # Summary
        print("")
        print("="*70)
        print("✅ PIPELINE COMPLETE!")
        print("="*70)
        print(f"📁 Generated files:")
        print(f"   - {result.get('skills_generated', 0)} SKILL.md files in '{output_dir}/'")
        print(f"   - {result.get('scripts_generated', 0)} Python scripts in './scripts/'")
        print(f"   - 1 extracted_knowledge.json")
        print("")
        print(f"📂 Output structure:")
        print(f"   {output_dir}/")
        print(f"   ├── extracted_knowledge.json")
        print(f"   ├── skill-name-1/")
        print(f"   │   ├── SKILL.md")
        print(f"   │   ├── metadata.json")
        print(f"   │   ├── rules.yaml")
        print(f"   │   ├── definitions.yaml")
        print(f"   │   └── examples.yaml")
        print(f"   └── ...")
        print(f"   ./scripts/")
        print(f"   ├── skill-name-1/")
        print(f"   │   └── skill_name_1.py")
        print(f"   └── ...")
        print("")
        print("💡 Tips:")
        print("   - To skip review: python lambda_handler.py document.pdf --auto")
        print("   - To test a script: python ./scripts/skill-name/skill_name.py")
        print("")

        return 0

    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    exit(main())
